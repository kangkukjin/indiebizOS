"""
slide_ai.py - 강의 슬라이드 생성 전용 가벼운 AI

설계 원칙 (Step 3):
- 시스템 AI의 API 키·모델 설정만 빌려서 새 provider 인스턴스 생성 (싱글턴 캐시)
- 인지 파이프라인(의식+평가) 우회 — 한 번의 process_message 호출로 끝
- 도구 없음(tools=[]) — JSON 스펙만 생성
- 컨텍스트는 호출자가 명시 조립: 강의 메타 + 재료 색인 + 데크 개요 + (편집 모드면) focus 슬라이드 + 누적 메모
- 채팅 히스토리 안 줌 — 매 호출이 self-contained, 22턴 컨텍스트 재구축 문제 회피
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

# backend 경로를 sys.path에 추가 (system_ai_core, providers 임포트용)
_THIS_DIR = Path(__file__).resolve().parent
_BACKEND_PATH = _THIS_DIR.parents[4] / "backend"
if str(_BACKEND_PATH) not in sys.path:
    sys.path.insert(0, str(_BACKEND_PATH))


# ─────────────────────────────────────────────────────────────────────
# 가벼운 provider 싱글턴
# ─────────────────────────────────────────────────────────────────────

_slide_ai_provider = None
_slide_ai_config_sig = None


def _load_system_ai_config() -> dict:
    """시스템 AI 설정 로드 (system_ai_core 우회 — 가볍게)."""
    from runtime_utils import get_base_path
    cfg_path = get_base_path() / "data" / "system_ai_config.json"
    if cfg_path.exists():
        with open(cfg_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"provider": "anthropic", "model": "claude-sonnet-4-20250514", "apiKey": ""}


def _get_slide_ai():
    """슬라이드 생성 전용 provider 반환. 시스템 AI 설정 변경 시 자동 재생성."""
    global _slide_ai_provider, _slide_ai_config_sig

    cfg = _load_system_ai_config()
    api_key = cfg.get("apiKey", "").strip()
    if not api_key:
        raise RuntimeError(
            "시스템 AI API 키가 설정되지 않았습니다. "
            "설정 → 시스템 AI에서 API 키를 입력하세요."
        )

    provider_name = cfg.get("provider", "anthropic").strip()
    model_name = cfg.get("model", "").strip()
    sig = (provider_name, model_name, api_key[-8:])  # 키 변경 감지용

    if _slide_ai_provider is not None and _slide_ai_config_sig == sig:
        return _slide_ai_provider

    from providers import get_provider
    provider = get_provider(
        provider_name,
        api_key=api_key,
        model=model_name,
        system_prompt=_SYSTEM_PROMPT,
        tools=[],
    )
    provider.init_client()
    _slide_ai_provider = provider
    _slide_ai_config_sig = sig
    print(f"[SlideAI] 초기화 완료 ({provider_name}/{model_name})")
    return provider


def reset_slide_ai():
    """설정 변경 시 캐시 초기화."""
    global _slide_ai_provider, _slide_ai_config_sig
    _slide_ai_provider = None
    _slide_ai_config_sig = None


# ─────────────────────────────────────────────────────────────────────
# 시스템 프롬프트
# ─────────────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """당신은 강의 슬라이드 생성 전문 AI다.

# 역할
강의자(사용자)가 한 장씩 협업으로 강의 슬라이드를 빚어간다. 매 호출마다 당신은 **단 한 장의 슬라이드 JSON 스펙**을 생성한다.

# 우선순위 (절대 규칙)
1. **강의자의 현재 요청이 최우선**. 누적 메모나 직전 슬라이드의 흐름보다 우선한다.
2. 강의자가 "X에 대한 슬라이드"라고 명시하면, 그 X를 책 본문에서 직접 찾아 명제로 추출. 첫 장의 메타포·톤에 묶여 X를 빗겨가지 말 것.
3. **책 본문 인용 의무**: 강의 재료(책 docx 등)가 컨텍스트에 있으면 거기에서 명제·표현·구체 사례를 가져와라. 당신의 일반 지식에서 만들어내지 말 것.
4. 누적 메모는 **톤·메타포 일관성** 참고용일 뿐 — 다음 슬라이드의 주제를 결정하지 않는다.

# 핵심 원칙 (lecture_slide_principles 요약)
1. **제목 = 명제** — 청중이 제목만 봐도 핵심을 얻을 수 있는 단정문. 챕터명·라벨 금지.
   - ✗ "프롬프트에서 에이전틱 루프로" (라벨)
   - ✓ "한 번의 속삭임에서 끊임없는 대화로" (명제)
2. **한 슬라이드 = 한 아이디어** — 두 가지를 가르치려면 두 슬라이드.
3. **본문은 키워드 단위** — 책 본문을 그대로 옮기지 말 것. body 1~2문장, bullets 3~5개.
4. **수치·고유명사는 시각 요소로 격리** — 평문에 묻지 말고 별도 박스.
5. **누적 메모 반영** — 강의자가 이미 거부한 톤·표현은 다시 쓰지 말 것. 채택한 메타포·강조 표현은 이어갈 것.

# 슬라이드 5종 수사 장치 → 레이아웃
| 명제 유형 | 추천 layout (텍스트형) | 추천 layout (일러스트형) |
|----------|---------------------|------------------------|
| 표지 | hero | hero_illustration |
| 두 항 대비 (A vs B) | comparison_table | split_concept / comparison_iconic |
| 단계·흐름·과정 | lecture_body (bullets) | illustration_overlay / illustration_anchor |
| 메타포·은유 | metaphor_story | illustration_anchor / illustration_background |
| 수치·팩트 강조 | factbox | comparison_iconic |
| 핵심 인용·회수 | quote | illustration_background |
| 동심원·계층 구조 | — | illustration_overlay (동심원/계층) |

# 일러스트 레이아웃 사용 가이드
다음 레이아웃들은 일러스트가 **반드시 필요**하다 — 선택 시 응답에 `illustrations` 필드도 같이 넣어야 한다:

- **hero_illustration**: 중앙 영웅 일러스트 + 상하 텍스트 (표지·부 도입)
- **illustration_anchor**: 상단 큰 일러스트 + 하단 텍스트(title/body/takeaway). 개념 설명
- **split_concept**: 좌우 대비. **좌우 각각** 일러스트 필요. left_title/left_body, right_title/right_body, conclusion
- **illustration_background**: 전면 배경 일러스트 + 오버레이 텍스트. text_align: "left"/"center"
- **illustration_overlay**: 풀-블리드 일러스트 + 자유 좌표 한글 라벨 박스. NotebookLM 양식의 다이어그램형
- **comparison_iconic**: 이모지/일러스트 헤더 비교표 — columns[{title,subtitle,icon,highlighted}], rows[{label, cells:[...]}]. 단순 표가 아니라 시각적 비교

일러스트 선택 기준:
- 메타포·인과 흐름·구조 다이어그램·강한 시각 임팩트가 필요한 명제 → 일러스트 layout
- 평범한 정의·본문·표 → 텍스트 layout (속도 빠름)
- **이미지 API 비용**이 들기 때문에 매 슬라이드 일러스트 쓰지 말 것. 핵심 명제·표지·메타포·구조 정리 슬라이드 위주.

# 자유형 레이아웃 (custom) — 틀에서 벗어나야 할 때
위의 고정 레이아웃 중 어느 것도 명제를 제대로 담지 못할 때, `layout: "custom"`을 선택하고 슬라이드 HTML을 **직접 작성**한다. 12개 틀에 갇히지 말고 내용에 맞는 배치를 새로 디자인하라.

**언제 쓰나**
- 타임라인·플로우차트·계층 트리·매트릭스·동심원 등 고정 틀에 없는 구조
- 여러 요소를 자유 좌표로 배치해야 하는 다이어그램
- 강의자가 "자유롭게/창의적으로/이 틀 말고" 라고 명시할 때
- 평범한 본문·표·표지는 **고정 레이아웃이 더 빠르고 안정적** — 굳이 custom 쓰지 말 것

**디자인 시스템 자동 상속 (반드시 지킬 것)**
custom HTML은 선택된 디자인 시스템(vintage_book 등)의 CSS 래퍼 안에 삽입된다. 색·폰트가 자동으로 흐르게 하려면 **하드코딩 색상 금지**, 반드시 CSS 변수를 써라:
- 글자색: `hsl(var(--foreground))` / 보조: `hsl(var(--muted-foreground))`
- 강조색: `hsl(var(--accent))` / 배경: `hsl(var(--background))` / 면: `hsl(var(--muted))`
- 테두리: `hsl(var(--border))`
- 사용 가능 자산: **Tailwind 유틸리티 클래스**(cdn), **animate.css** 클래스, **lucide 아이콘**(`<i data-lucide="아이콘명"></i>`), 이모지, **인라인 SVG**.
- 시각 요소는 SVG/이모지/아이콘/CSS로 직접 그려라 — **custom은 이미지 API를 호출하지 않는다**(빠르고 무료). `illustrations` 필드 쓰지 말 것.

**작성 규칙**
- 최상위 요소는 슬라이드 전체(1280×720)를 채운다: `<div class="w-full h-full ...">` 로 시작.
- 한글 텍스트는 HTML이므로 그대로 써도 된다(이미지가 아님).
- 메시지 큐레이션 원칙은 그대로 — 제목=명제, 한 슬라이드 한 아이디어, 키워드 단위.

## 자유형 layout 예시 (custom)
```json
{
  "slide": {
    "layout": "custom",
    "title": "에이전트 인지 루프의 4단계",
    "custom_html": "<div class=\\"w-full h-full flex flex-col p-16\\" style=\\"color: hsl(var(--foreground))\\"><p class=\\"text-sm font-semibold uppercase tracking-wider mb-2\\" style=\\"color: hsl(var(--accent))\\">2부 · 구조</p><h1 class=\\"text-4xl font-black mb-10\\">생각은 한 바퀴 돈다</h1><div class=\\"flex items-center justify-between flex-1\\"><div class=\\"flex-1 text-center\\"><div class=\\"text-5xl mb-3\\">🧭</div><div class=\\"font-bold text-xl\\">분류</div><div class=\\"text-sm\\" style=\\"color: hsl(var(--muted-foreground))\\">반사</div></div><i data-lucide=\\"arrow-right\\" style=\\"color: hsl(var(--accent))\\"></i><div class=\\"flex-1 text-center\\"><div class=\\"text-5xl mb-3\\">💭</div><div class=\\"font-bold text-xl\\">의식</div><div class=\\"text-sm\\" style=\\"color: hsl(var(--muted-foreground))\\">계획</div></div><i data-lucide=\\"arrow-right\\" style=\\"color: hsl(var(--accent))\\"></i><div class=\\"flex-1 text-center\\"><div class=\\"text-5xl mb-3\\">⚙️</div><div class=\\"font-bold text-xl\\">실행</div></div><i data-lucide=\\"arrow-right\\" style=\\"color: hsl(var(--accent))\\"></i><div class=\\"flex-1 text-center\\"><div class=\\"text-5xl mb-3\\">🔍</div><div class=\\"font-bold text-xl\\">평가</div><div class=\\"text-sm\\" style=\\"color: hsl(var(--muted-foreground))\\">성찰</div></div></div></div>"
  },
  "reasoning": "4단계 순환은 고정 틀에 없어 자유형으로 가로 플로우를 직접 배치",
  "speaker_note": "..."
}
```
custom은 `title`(데크 목록 표시용)과 `custom_html` 두 필드만 있으면 된다. 다른 레이아웃 필드는 생략.

# 출력 형식
**반드시 다음 JSON 한 객체만 반환**. 코드 펜스 ```json ... ``` 안에 넣어도 되고, 그냥 객체만 줘도 된다. 다른 설명 텍스트 금지.

## 텍스트 layout 예시 (일러스트 없음)
```json
{
  "slide": {
    "layout": "lecture_body",
    "eyebrow": "1부 · 정의",
    "title": "AI 모델은 AI의 실체가 아니다",
    "body": "AI 모델 자체는 정적인 지식 덩어리. 실제로 일을 하려면 하네스가 필요하다.",
    "bullets": ["모델 = 학습된 지식", "하네스 = 도구·기억·환경 연결", "둘이 결합해야 행동 가능"],
    "footer": "1/15"
  },
  "reasoning": "...",
  "speaker_note": "...",
  "memo_signals": { "metaphors_adopted": ["하네스 = 도구·기억·환경 연결"] }
}
```

## 일러스트 layout 예시 (split_concept) — illustrations 필드 필수
```json
{
  "slide": {
    "layout": "split_concept",
    "eyebrow": "1부 · 대비",
    "left_title": "유리병 속의 뇌",
    "left_body": "AI 모델 자체. 지식은 있지만 손이 없다.",
    "right_title": "깨진 유리병 + 하네스",
    "right_body": "모델이 도구·기억·환경과 연결된 상태. 비로소 움직인다.",
    "conclusion": "모델 + 하네스 = 실제 운동"
  },
  "illustrations": {
    "left_image_prompt": "A human brain enclosed in a transparent glass jar on a wooden table, isolated, no connection to outside, hand-drawn ink illustration style",
    "right_image_prompt": "The same brain emerging from a cracked glass jar, with mechanical arms and wires connecting to surrounding objects (books, tools, a window showing the world), dynamic and active composition"
  },
  "reasoning": "..."
}
```

## 일러스트 layout 예시 (illustration_anchor 또는 hero_illustration) — 단일 이미지
```json
{
  "slide": {
    "layout": "illustration_anchor",
    "eyebrow": "1부 · 메타포",
    "title": "햄릿의 배우와 연출가",
    "body": "한 명의 배우가 연출가의 한마디에 따라 다른 햄릿이 된다.",
    "takeaway": "AI 모델 = 배우, 시스템 프롬프트 = 연출가"
  },
  "illustrations": {
    "image_prompt": "A theatrical scene with an actor in Hamlet costume standing center stage, a director gesturing from the side with a single word in a speech bubble, vintage theatrical illustration style"
  }
}
```

# 이미지 프롬프트 작성 규칙 (illustrations 필드)
1. **영어로 작성**. 한국어 프롬프트는 이미지 생성기가 잘 못 알아듣는다.
2. **다이어그램·인포그래픽 지향**. "장면 묘사"가 아니라 "구조·관계·은유의 시각화".
3. **간결**. 1~3문장. 디자인 시스템 톤(vintage_book / sf_blueprint 등)은 자동 적용되니 다시 쓰지 말 것.
4. **금지**: 한글/Hangul 텍스트가 이미지 안에 들어가지 않게 하라 (영문 라벨은 OK).
5. **split_concept**: 좌우가 **시각적으로 명확히 대비**되어야 — 같은 주제, 다른 상태.

# 각 layout의 필드
## 텍스트 layout (일러스트 없음)
- **hero**: title, subtitle, eyebrow
- **lecture_body**: eyebrow, title, body, bullets[], quote, footer
- **metaphor_story**: eyebrow, title, label, story, takeaway
- **comparison_table**: eyebrow, title, headers[], rows[[...]]
- **factbox**: eyebrow, title, body, items[], source
- **quote**: quote, attribution, context

## 일러스트 layout (illustrations 필드 필수)
- **hero_illustration**: title, subtitle, eyebrow + illustrations.image_prompt
- **illustration_anchor**: eyebrow, title, body, takeaway + illustrations.image_prompt
- **illustration_background**: eyebrow, title, subtitle, text_align("left"|"center") + illustrations.image_prompt
- **split_concept**: eyebrow, left_title, left_body, right_title, right_body, conclusion + illustrations.left_image_prompt + illustrations.right_image_prompt
- **illustration_overlay**: eyebrow, title + illustrations.image_prompt (자유 좌표 라벨 박스는 별도 — 본 MVP에서는 단순 오버레이만)
- **comparison_iconic**: eyebrow, title, label_header, columns[{title,subtitle,icon,highlighted}], rows[{label, cells:[...]}] (icon은 이모지 권장, illustrations 필요 없음)

## 자유형 layout (틀에서 벗어날 때)
- **custom**: title(데크 표시용) + custom_html(슬라이드 전체 HTML). 디자인 시스템 CSS 변수 필수, illustrations 없음. 위 "자유형 레이아웃" 섹션 참조.

필요 없는 필드는 생략. 항상 eyebrow에 "부·장·라벨" 형식으로 위치 표시 (예: "1부 · 정의").

# 협업 흐름의 자각
- 첫 호출이면: 강의자의 의도를 짚어 표지·개요부터 차근차근.
- 이미 슬라이드가 있으면: 데크 개요를 보고 "다음 한 장"을 자연스럽게.
- 편집 요청이면: focus 슬라이드의 현재 spec을 받아 강의자 코멘트만 반영. 다른 부분은 보존.
- 강의자가 짧게 "다음" 류로만 말하면 데크 흐름에 따라 자연스러운 다음 명제를 제안. 거대한 outline은 강요하지 말 것.

# 금지
- ❌ 채팅 히스토리 인용 ("앞서 말씀하신...") — 매 호출이 self-contained
- ❌ 여러 슬라이드 동시 생성 — 항상 한 장
- ❌ JSON 외의 자유 텍스트 출력
- ❌ 강의자가 거부한 톤 재사용 (누적 메모의 tone_rejected 확인)
"""


# ─────────────────────────────────────────────────────────────────────
# 컨텍스트 조립
# ─────────────────────────────────────────────────────────────────────

_MATERIAL_INLINE_BYTES = 4000  # 재료당 inline 최대 바이트
_MATERIAL_TOTAL_BYTES = 16000  # 전체 재료 inline 합계 상한


def _read_text_safe(path: Path, limit: int) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception as e:
        return f"(읽기 실패: {e})"


def _read_docx_safe(path: Path, limit_chars: int = 4000) -> str:
    """python-docx로 docx 텍스트 추출 (상한 적용)."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "(python-docx 미설치 — 파일명만 인식)"
    try:
        doc = Document(str(path))
        out = []
        total = 0
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            out.append(t)
            total += len(t)
            if total >= limit_chars:
                out.append(f"... (이후 생략, 총 {len(doc.paragraphs)}문단)")
                break
        return "\n".join(out)
    except Exception as e:
        return f"(docx 읽기 실패: {e})"


def _build_materials_block(deck: dict, lecture_dir: Path) -> str:
    """재료 목록 + 일부 inline 텍스트."""
    materials = deck.get("materials", [])
    if not materials:
        return "(재료 없음 — 강의자가 직접 명제·내용을 알려줄 것)"

    lines = ["## 재료 목록"]
    inlined = []
    total = 0
    for m in materials:
        rel = m.get("file", "")
        mtype = m.get("type", "?")
        lines.append(f"- {rel} [{mtype}]")
        if total >= _MATERIAL_TOTAL_BYTES:
            continue

        path = lecture_dir / rel
        if not path.exists():
            continue

        chunk = None
        if mtype in ("text", "md", "txt"):
            chunk = _read_text_safe(path, _MATERIAL_INLINE_BYTES)
        elif mtype == "docx":
            chunk = _read_docx_safe(path, _MATERIAL_INLINE_BYTES)
        # pdf, image는 inline 안 함 (파일명만)

        if chunk:
            header = f"\n### {rel}\n"
            available = _MATERIAL_TOTAL_BYTES - total - len(header)
            if available < 200:
                continue
            chunk = chunk[:available]
            inlined.append(header + chunk)
            total += len(header) + len(chunk)

    block = "\n".join(lines)
    if inlined:
        block += "\n\n## 재료 본문 (일부 inline)\n" + "\n".join(inlined)
    else:
        block += "\n\n(재료 본문은 inline 안 됨 — 강의자에게 핵심 내용 확인 필요)"
    return block


def _build_deck_overview(deck: dict) -> str:
    order = deck.get("slide_order", [])
    if not order:
        return "(아직 슬라이드 없음 — 첫 슬라이드 생성)"
    slides = deck.get("slides", {})
    lines = ["## 현재 데크 (이미 만들어진 슬라이드)"]
    for i, sid in enumerate(order, 1):
        s = slides.get(sid, {})
        lines.append(f"{i}. [{s.get('layout', '?')}] {s.get('title', '(제목 없음)')}")
    return "\n".join(lines)


def _build_memo_block(deck: dict) -> str:
    memo = deck.get("cumulative_memo", {})
    parts = []
    if memo.get("tone_preferred"):
        parts.append(f"### 강의자 선호 톤·표현\n" + "\n".join(f"- {x}" for x in memo["tone_preferred"]))
    if memo.get("tone_rejected"):
        parts.append(f"### 거부한 방향 (다시 쓰지 말 것)\n" + "\n".join(f"- {x}" for x in memo["tone_rejected"]))
    if memo.get("metaphors_adopted"):
        parts.append(f"### 채택한 메타포\n" + "\n".join(f"- {x}" for x in memo["metaphors_adopted"]))
    if memo.get("decisions"):
        parts.append(f"### 누적 결정사항\n" + "\n".join(f"- {x}" for x in memo["decisions"]))
    if not parts:
        return "(누적 메모 없음 — 첫 협업 라운드)"
    return "\n\n".join(parts)


def _build_neighbor_block(neighbor_briefs: Optional[list]) -> str:
    """앞뒤 이웃 슬라이드 요약 — 텍스트 경로에서 톤·흐름 일관성 참고용(주제 결정 아님)."""
    if not neighbor_briefs:
        return ""
    lines = ["## 앞뒤 이웃 슬라이드 (스타일·흐름 참고용 — 내용 복제 금지)"]
    for b in neighbor_briefs:
        pos = b.get("position", "")
        lines.append(f"- [{pos}] [{b.get('layout', '?')}] {b.get('title', '(제목 없음)')}")
    lines.append(
        "→ 이 슬라이드가 같은 덱의 일부로 보이게 톤·강조·layout 언어를 맞추되, "
        "이웃의 제목·내용을 베끼지 말고 이번 명제만 담아라."
    )
    return "\n".join(lines)


def _build_meta_block(deck: dict) -> str:
    return (
        f"## 강의 메타\n"
        f"- 제목: {deck.get('title', '')}\n"
        f"- 청중: {deck.get('audience') or '(미지정)'}\n"
        f"- 한 줄 요지: {deck.get('thesis') or '(미지정)'}\n"
        f"- 분량: {deck.get('duration_minutes') or 0}분\n"
        f"- 디자인: {deck.get('design_system', 'vintage_book')}"
    )


def build_prompt(
    deck: dict,
    lecture_dir: Path,
    user_instruction: str,
    focus_slide: Optional[dict] = None,
    forced_layout: Optional[str] = None,
    neighbor_briefs: Optional[list] = None,
) -> str:
    """전체 사용자 프롬프트 조립.

    핵심 순서: 강의자의 요청 → 컨텍스트(메타·데크·메모) → 출력 지시.
    instruction을 최상단에 두는 이유 — LM의 attention이 앞쪽에 강하고, 누적 메모가
    과도하게 영향을 미쳐서 강의자 instruction을 무시하는 사례가 관찰됨.

    ★자료(원고 등)는 더 이상 매 슬라이드 프롬프트에 넣지 않는다(2026-06-24). 자료는
    '초안 생성'(outline_from_materials)에서만 통째로 읽는다. 한 장 생성은 instruction이
    이끈다 — 일괄 생성이면 초안 단계가 자료를 소화해 instruction에 담아 넘긴다.
    """
    blocks = [
        _build_instruction_block(user_instruction, focus_slide, forced_layout),
        _build_meta_block(deck),
        _build_deck_overview(deck),
        f"## 강의자의 호흡 (누적 메모)\n{_build_memo_block(deck)}",
    ]
    nb = _build_neighbor_block(neighbor_briefs)
    if nb:
        blocks.append(nb)
    if focus_slide:
        focus_json = json.dumps(focus_slide, ensure_ascii=False, indent=2)
        blocks.append(
            f"## 현재 슬라이드 (편집 대상)\n```json\n{focus_json}\n```\n"
            f"(layout은 가급적 보존. 강의자가 명시적으로 바꾸지 않은 필드는 그대로 유지.)"
        )
    blocks.append(_OUTPUT_DIRECTIVE)
    return "\n\n".join(blocks)


def _build_dynamic_prompt(
    deck: dict,
    user_instruction: str,
    focus_slide: Optional[dict] = None,
    forced_layout: Optional[str] = None,
    neighbor_briefs: Optional[list] = None,
) -> str:
    """캐시 위에 얹을 **동적** 부분만 — 책 본문은 캐시에 있다는 전제.

    캐시 콘텐츠: 시스템 프롬프트 + 강의 재료 전체.
    동적: 강의자 instruction(최상단) + 강의 메타 + 데크 개요 + 누적 메모 + (편집 시 focus slide).
    """
    blocks = [
        _build_instruction_block(user_instruction, focus_slide, forced_layout),
        (
            "## 강의 재료 위치\n"
            "(이 대화의 캐시된 컨텍스트에 강의 재료(책 본문 전체)가 포함되어 있다. "
            "**반드시 그 책 본문에서 명제·인용·구체 표현을 직접 추출하라** — "
            "당신이 기억하는 일반 지식이 아니라 강의자의 책에 있는 그대로.)"
        ),
        _build_meta_block(deck),
        _build_deck_overview(deck),
        f"## 강의자의 호흡 (누적 메모)\n{_build_memo_block(deck)}",
    ]
    nb = _build_neighbor_block(neighbor_briefs)
    if nb:
        blocks.append(nb)
    if focus_slide:
        focus_json = json.dumps(focus_slide, ensure_ascii=False, indent=2)
        blocks.append(
            f"## 현재 슬라이드 (편집 대상)\n```json\n{focus_json}\n```\n"
            f"(layout은 가급적 보존. 강의자가 명시적으로 바꾸지 않은 필드는 그대로 유지.)"
        )
    blocks.append(_OUTPUT_DIRECTIVE)
    return "\n\n".join(blocks)


def _build_instruction_block(
    user_instruction: str,
    focus_slide: Optional[dict],
    forced_layout: Optional[str] = None,
) -> str:
    """프롬프트 최상단의 강의자 요청 — 최우선 신호.

    forced_layout이 있으면 강한 강제 블록 추가 — UI에서 명시적으로 layout을 선택한 경우.
    """
    mode = "편집(재생성)" if focus_slide else "신규 슬라이드 생성"
    parts = [
        f"# 강의자의 요청 ({mode}) — 최우선 신호\n",
        f"> {user_instruction}\n",
        (
            "이 요청은 누적 메모나 직전 슬라이드의 흐름보다 **우선**한다. "
            "강의자가 명시한 주제·각도·표현이 있으면 그것을 정확히 반영하라. "
            "누적 메모는 톤·메타포 일관성용 참고일 뿐 주제를 결정하지 않는다."
        ),
    ]
    if forced_layout == "custom":
        parts.append(
            "\n## ⚠ Layout 강제 지정 — 자유형(custom)\n"
            "강의자가 UI에서 **자유형(custom)**을 명시적으로 지정했다. "
            "고정 레이아웃을 쓰지 말고 `layout: \"custom\"` + `custom_html`로 슬라이드 HTML을 직접 작성하라. "
            "디자인 시스템 CSS 변수(`hsl(var(--foreground))`, `hsl(var(--accent))` 등)와 Tailwind/이모지/SVG/lucide만 사용. "
            "`illustrations` 필드는 쓰지 말 것(이미지 API 호출 안 함). "
            "자세한 규칙은 시스템 프롬프트의 '자유형 레이아웃' 섹션을 따른다."
        )
    elif forced_layout:
        parts.append(
            f"\n## ⚠ Layout 강제 지정\n"
            f"강의자가 UI에서 **layout = `{forced_layout}`**을 명시적으로 지정했다. "
            f"반드시 이 layout으로 슬라이드를 생성하라 — 다른 layout 선택 금지. "
            f"layout 필드 값은 `{forced_layout}` 그대로 출력. "
            f"이 layout이 일러스트가 필요한 종류(`hero_illustration`, `illustration_anchor`, "
            f"`illustration_background`, `split_concept`, `illustration_overlay`)면 "
            f"`illustrations` 필드도 누락하지 말 것."
        )
    return "\n".join(parts)


_OUTPUT_DIRECTIVE = (
    "# 출력\n"
    "위 정보를 바탕으로 **단 한 장의 슬라이드 JSON**을 출력하라.\n"
    "- 강의자의 요청이 구체적 주제를 지목하면 그 주제로\n"
    "- 책 본문에서 직접 명제·인용을 추출 (외부 일반 지식 X)\n"
    "- 일러스트 layout을 고르면 `illustrations` 필드 누락 금지"
)


# ─────────────────────────────────────────────────────────────────────
# 책 본문 전체 텍스트 — 캐시 콘텐츠용 (상한 없음)
# ─────────────────────────────────────────────────────────────────────

_FULL_TEXT_MAX_BYTES = 800_000  # ~800KB 텍스트 (gemini-3.5-flash 1M 토큰 컨텍스트 안에 충분)


def _read_text_full(path: Path, limit: int) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception as e:
        return f"(읽기 실패: {e})"


def _read_docx_full(path: Path, limit_chars: int) -> str:
    """python-docx로 docx 텍스트 전체 추출 (상한 적용)."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "(python-docx 미설치)"
    try:
        doc = Document(str(path))
        out = []
        total = 0
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            out.append(t)
            total += len(t)
            if total >= limit_chars:
                out.append(f"... (이후 생략, 총 {len(doc.paragraphs)}문단)")
                break
        return "\n".join(out)
    except Exception as e:
        return f"(docx 읽기 실패: {e})"


def _build_full_materials_text(deck: dict, lecture_dir: Path) -> str:
    """캐시에 올릴 재료 본문 — 책 한 권 전체. _FULL_TEXT_MAX_BYTES 상한."""
    materials = deck.get("materials", [])
    if not materials:
        return ""

    parts = []
    total = 0
    for m in materials:
        rel = m.get("file", "")
        mtype = m.get("type", "?")
        path = lecture_dir / rel
        if not path.exists():
            continue

        remaining = _FULL_TEXT_MAX_BYTES - total
        if remaining < 1000:
            break

        if mtype in ("text", "md", "txt"):
            chunk = _read_text_full(path, remaining)
        elif mtype == "docx":
            chunk = _read_docx_full(path, remaining)
        else:
            continue  # pdf/image는 일단 스킵 (Step 5b에서 처리)

        if chunk:
            header = f"\n## 재료: {rel}\n\n"
            parts.append(header + chunk)
            total += len(header) + len(chunk)

    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────
# 강의별 캐시 매니저 (explicit prompt caching)
# ─────────────────────────────────────────────────────────────────────
#
# Gemini의 explicit cached content API를 강의별로 사용한다.
# 캐시 콘텐츠 = 시스템 프롬프트 + 책 본문 (둘 다 호출 간 고정).
# 동적 부분 = 강의 메타 + 데크 개요 + 누적 메모 + 강의자 instruction.
#
# 무효화:
#   - 재료 추가/삭제 → invalidate_lecture_cache(lecture_id)
#   - 시스템 프롬프트 변경 (서명 불일치) → 자동 재생성
#   - TTL 만료 (1시간) → 자동 재생성

_CACHE_TTL_SECONDS = 3600

# 캐시 최소 크기 — Gemini는 일정 토큰 미만이면 cache 거부. 한글 ~15k자 미만이면 그냥 inline.
_CACHE_MIN_MATERIAL_BYTES = 12_000


def _compute_cache_signature(deck: dict, lecture_dir: Path, model_name: str) -> str:
    """캐시 무효화 키 — 재료·시스템 프롬프트·모델 변경 감지."""
    h = hashlib.sha256()
    h.update(_SYSTEM_PROMPT.encode("utf-8"))
    h.update(f"|model={model_name}|".encode())
    for m in deck.get("materials", []):
        rel = m.get("file", "")
        p = lecture_dir / rel
        if p.exists():
            stat = p.stat()
            h.update(f"|{rel}|mtime={stat.st_mtime}|size={stat.st_size}".encode())
    return h.hexdigest()[:16]


def _cache_state_path(lecture_id: str) -> Path:
    """캐시 메타 저장 위치. lecture_dir/_slide_cache_state.json."""
    from runtime_utils import get_base_path
    return get_base_path() / "outputs" / "lectures" / lecture_id / "_slide_cache_state.json"


def _load_cache_state(lecture_id: str) -> Optional[dict]:
    p = _cache_state_path(lecture_id)
    if not p.exists():
        return None
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _save_cache_state(lecture_id: str, state: dict) -> None:
    p = _cache_state_path(lecture_id)
    p.parent.mkdir(parents=True, exist_ok=True)
    with open(p, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


def _delete_cache_state(lecture_id: str) -> None:
    p = _cache_state_path(lecture_id)
    if p.exists():
        try:
            p.unlink()
        except Exception:
            pass


def _get_genai_client():
    """google-genai Client 인스턴스 (싱글턴). cache API 호출용."""
    cfg = _load_system_ai_config()
    api_key = cfg.get("apiKey", "").strip()
    if not api_key:
        raise RuntimeError("시스템 AI API 키가 없습니다.")
    from google import genai
    return genai.Client(api_key=api_key)


def invalidate_lecture_cache(lecture_id: str) -> None:
    """강의의 캐시 삭제 — 재료 변경 시 핸들러에서 호출.

    Gemini 쪽 캐시 + 로컬 state 모두 정리.
    """
    state = _load_cache_state(lecture_id)
    if state and state.get("cache_name"):
        try:
            client = _get_genai_client()
            client.caches.delete(name=state["cache_name"])
            print(f"[SlideCache] 삭제: {state['cache_name']} (lecture={lecture_id})")
        except Exception as e:
            print(f"[SlideCache] 삭제 실패 (무시): {e}")
    _delete_cache_state(lecture_id)


def _ensure_lecture_cache(deck: dict, lecture_dir: Path) -> Optional[str]:
    """강의별 캐시 보장. 캐시 이름 반환. 캐시 못 만들면 None (호출자가 폴백 처리)."""
    lecture_id = deck["lecture_id"]
    cfg = _load_system_ai_config()
    model_name = cfg.get("model", "").strip()
    if not cfg.get("apiKey", "").strip():
        return None

    signature = _compute_cache_signature(deck, lecture_dir, model_name)

    # 기존 캐시 유효성 검사
    state = _load_cache_state(lecture_id)
    if state and state.get("signature") == signature:
        # 서명 같음 → 캐시 그대로 쓸 수 있을지 확인 (TTL 만료 가능)
        cache_name = state.get("cache_name")
        if cache_name:
            try:
                client = _get_genai_client()
                client.caches.get(name=cache_name)
                return cache_name  # 살아있음
            except Exception:
                # TTL 만료 또는 삭제됨 — 재생성
                _delete_cache_state(lecture_id)

    # 새 캐시 생성 — 재료 크기 확인
    materials_text = _build_full_materials_text(deck, lecture_dir)
    if not materials_text or len(materials_text.encode("utf-8")) < _CACHE_MIN_MATERIAL_BYTES:
        # 재료가 짧으면 캐시 안 만들고 그냥 inline 폴백
        return None

    # 기존 캐시(서명 다름)가 있으면 미리 삭제
    if state and state.get("cache_name"):
        try:
            client = _get_genai_client()
            client.caches.delete(name=state["cache_name"])
        except Exception:
            pass
        _delete_cache_state(lecture_id)

    try:
        from google import genai
        from google.genai import types
        client = _get_genai_client()
        import time
        t0 = time.time()
        cache = client.caches.create(
            model=model_name,
            config=types.CreateCachedContentConfig(
                displayName=f"lecture-{lecture_id}",
                systemInstruction=_SYSTEM_PROMPT,
                contents=[{
                    "role": "user",
                    "parts": [{"text": (
                        "# 강의 재료 (책 본문 등) — 이후의 모든 슬라이드 생성에서 참조하라.\n\n"
                        + materials_text
                    )}],
                }],
                ttl=f"{_CACHE_TTL_SECONDS}s",
            ),
        )
        elapsed = time.time() - t0
        _save_cache_state(lecture_id, {
            "cache_name": cache.name,
            "signature": signature,
            "model": model_name,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "ttl_seconds": _CACHE_TTL_SECONDS,
            "materials_bytes": len(materials_text.encode("utf-8")),
        })
        print(
            f"[SlideCache] 생성: {cache.name} (lecture={lecture_id}, "
            f"materials={len(materials_text):,}자, {elapsed:.1f}s, TTL=1h)"
        )
        return cache.name
    except Exception as e:
        print(f"[SlideCache] 생성 실패 (캐시 없이 진행): {e}")
        return None


def _generate_with_cache(
    cache_name: str,
    user_prompt: str,
    model_name: str,
) -> str:
    """캐시 참조로 단일 호출. 응답 text 반환."""
    from google.genai import types
    client = _get_genai_client()
    response = client.models.generate_content(
        model=model_name,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            cached_content=cache_name,
        ),
    )
    return response.text or ""


# ─────────────────────────────────────────────────────────────────────
# JSON 추출 (AI 응답에서)
# ─────────────────────────────────────────────────────────────────────

_JSON_FENCE_RE = re.compile(r"```(?:json)?\s*(\{.*?\})\s*```", re.DOTALL)


def extract_json(text: str) -> dict:
    """AI 응답 텍스트에서 JSON 객체 추출. 실패 시 ValueError.

    1. ```json ... ``` 펜스 우선
    2. 본문에서 첫 {부터 짝 맞는 }까지
    """
    # 1. 펜스
    m = _JSON_FENCE_RE.search(text)
    if m:
        candidate = m.group(1)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass  # 펜스 안인데 깨졌으면 fall through

    # 2. 균형 잡힌 {} 추출
    start = text.find('{')
    if start == -1:
        raise ValueError("AI 응답에 JSON이 없습니다.")
    depth = 0
    end = -1
    in_string = False
    escape = False
    for i in range(start, len(text)):
        ch = text[i]
        if escape:
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = i
                break
    if end == -1:
        raise ValueError("JSON 객체의 닫는 괄호를 찾지 못했습니다.")
    candidate = text[start:end + 1]
    try:
        return json.loads(candidate)
    except json.JSONDecodeError as e:
        raise ValueError(f"JSON 파싱 실패: {e}\n원문 일부: {candidate[:200]}")


# ─────────────────────────────────────────────────────────────────────
# 슬라이드 렌더링 (media_producer의 shadcn_slides 호출)
# ─────────────────────────────────────────────────────────────────────

def _load_shadcn_module():
    """media_producer/shadcn_slides.py를 동적 로드."""
    media_pkg = _THIS_DIR.parent / "media_producer"
    mod_path = media_pkg / "shadcn_slides.py"
    if not mod_path.exists():
        raise RuntimeError(f"shadcn_slides.py 없음: {mod_path}")

    if "shadcn_slides" in sys.modules:
        return sys.modules["shadcn_slides"]

    import importlib.util
    spec = importlib.util.spec_from_file_location("shadcn_slides", mod_path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules["shadcn_slides"] = mod
    spec.loader.exec_module(mod)
    return mod


def _load_media_handler():
    """media_producer 핸들러를 tool_loader 표준 경로로 로드.

    이전 버전은 importlib로 직접 로드했는데, exec_module이 실패해도 sys.modules에
    빈 모듈이 남아서 다음 호출에 attribute 누락 에러 발생. tool_loader는 캐싱 + 에러
    처리가 잘 돼있어서 안전.
    """
    from tool_loader import load_tool_handler
    handler = load_tool_handler("generate_gemini_image")
    if handler is None or not hasattr(handler, "execute"):
        raise RuntimeError(
            "generate_gemini_image 핸들러를 찾을 수 없음. "
            "media_producer 패키지가 설치돼있는지 확인하세요."
        )
    return handler


# ─────────────────────────────────────────────────────────────────────
# 일러스트 생성 — illustrations 필드 처리
# ─────────────────────────────────────────────────────────────────────

# illustrations 필드의 각 키 → slide_spec의 어떤 image_path에 주입할지
_ILLUSTRATION_KEY_MAP = {
    "image_prompt": "image_path",
    "left_image_prompt": "left_image_path",
    "right_image_prompt": "right_image_path",
    "avatar_prompt": "avatar_path",
}


def _generate_one_illustration(
    prompt: str,
    design_system: str,
    output_path: Path,
) -> str:
    """단일 이미지 생성. 성공 시 output_path의 절대경로 반환, 실패 시 예외.

    media_producer 핸들러의 표준 ToolContext 인터페이스로 호출:
    - 임시 project_path를 만들어 그 안의 outputs/에 생성 → 우리 위치로 이동
    - 이렇게 가야 ToolContext.output_dir() 규약과 충돌하지 않음
    """
    handler = _load_media_handler()

    from tool_context import ToolContext
    with tempfile.TemporaryDirectory() as tmp:
        ctx = ToolContext(project_path=tmp, tool_name="generate_gemini_image")
        tool_input = {
            "prompt": prompt,
            "style_preset": design_system,
            "aspect_ratio": "16:9",
            "image_size": "1K",
            "output_path": output_path.name,  # 파일명만 — ToolContext.output_dir()/{name}에 저장됨
        }
        result = handler.execute(tool_input, ctx)
        if isinstance(result, str) and ("오류" in result[:10] or "에러" in result[:10]):
            raise RuntimeError(f"이미지 생성 실패: {result[:300]}")

        # 결과는 {tmp}/outputs/{filename} 위치
        expected = Path(tmp) / "outputs" / output_path.name
        if not expected.exists():
            # 백업 탐색
            candidates = list(Path(tmp).rglob(output_path.name))
            if not candidates:
                raise RuntimeError(
                    f"이미지가 생성되지 않음 (예상 위치: {expected})\n"
                    f"핸들러 응답: {result[:300] if isinstance(result, str) else type(result).__name__}"
                )
            expected = candidates[0]

        output_path.parent.mkdir(parents=True, exist_ok=True)
        if output_path.exists():
            output_path.unlink()
        shutil.move(str(expected), str(output_path))
        return str(output_path.resolve())


def generate_slide_illustrations(
    illustrations: dict,
    design_system: str,
    slides_dir: Path,
    slide_id: str,
) -> dict:
    """illustrations dict의 각 prompt를 이미지로 생성. slide_spec에 주입할 키들 반환.

    예: {"image_prompt": "..."} → {"image_path": "/abs/path/s001_img.png"}
        {"left_image_prompt": "A", "right_image_prompt": "B"}
        → {"left_image_path": "...", "right_image_path": "..."}

    여러 이미지가 있으면 ThreadPoolExecutor로 병렬 생성 (image_gemini는 HTTP 호출).
    """
    if not illustrations:
        return {}

    # 생성할 작업 목록
    jobs = []
    for prompt_key, prompt_text in illustrations.items():
        if prompt_key not in _ILLUSTRATION_KEY_MAP:
            continue
        if not isinstance(prompt_text, str) or not prompt_text.strip():
            continue
        target_key = _ILLUSTRATION_KEY_MAP[prompt_key]
        # role 부분: "left_image_path" → "left", "image_path" → "img" 등
        role = prompt_key.replace("_image_prompt", "").replace("_prompt", "") or "img"
        filename = f"{slide_id}_{role}.png" if role != "img" else f"{slide_id}_img.png"
        output_path = slides_dir / filename
        jobs.append((target_key, prompt_text, output_path))

    if not jobs:
        return {}

    # 1개면 직렬, 여러 개면 병렬
    injected = {}
    if len(jobs) == 1:
        target_key, prompt_text, output_path = jobs[0]
        abs_path = _generate_one_illustration(prompt_text, design_system, output_path)
        injected[target_key] = abs_path
    else:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        with ThreadPoolExecutor(max_workers=len(jobs)) as ex:
            futures = {
                ex.submit(_generate_one_illustration, prompt_text, design_system, output_path): target_key
                for target_key, prompt_text, output_path in jobs
            }
            for fut in as_completed(futures):
                target_key = futures[fut]
                injected[target_key] = fut.result()  # 예외는 propagate
    return injected


def render_slide_to_files(
    spec: dict,
    design_system: str,
    slides_dir: Path,
    slide_id: str,
) -> dict:
    """슬라이드 spec → PNG/HTML 파일로 렌더. slides/{slide_id}.{png,html} 위치.

    Returns: {png_file: str (rel), html_file: str (rel), spec_file: str (rel)}
    """
    shadcn = _load_shadcn_module()

    with tempfile.TemporaryDirectory() as tmp:
        tool_input = {
            "slides": [spec],
            "design_system": design_system or "vintage_book",
            "output_dir": tmp,
        }
        result_str = shadcn.create_shadcn_slides(tool_input, tmp)
        # 결과는 JSON 문자열 또는 에러 문자열
        try:
            result = json.loads(result_str) if isinstance(result_str, str) else result_str
        except Exception:
            raise RuntimeError(f"slide_shadcn 응답 파싱 실패: {result_str[:500]}")

        if not result.get("success"):
            raise RuntimeError(f"slide_shadcn 실패: {result.get('message') or result}")

        images = result.get("images", [])
        htmls = result.get("html_files", [])
        if not images:
            raise RuntimeError("slide_shadcn이 이미지를 반환하지 않았습니다.")

        slides_dir.mkdir(parents=True, exist_ok=True)
        # 이전 파일이 있다면 덮어쓰기
        dest_png = slides_dir / f"{slide_id}.png"
        dest_html = slides_dir / f"{slide_id}.html"
        dest_spec = slides_dir / f"{slide_id}.json"

        if dest_png.exists():
            dest_png.unlink()
        if dest_html.exists():
            dest_html.unlink()

        shutil.move(images[0], dest_png)
        if htmls:
            try:
                shutil.move(htmls[0], dest_html)
            except Exception:
                pass

        with open(dest_spec, "w", encoding="utf-8") as f:
            json.dump(spec, f, ensure_ascii=False, indent=2)

    return {
        "png_file": f"slides/{slide_id}.png",
        "html_file": f"slides/{slide_id}.html",
        "spec_file": f"slides/{slide_id}.json",
    }


# ─────────────────────────────────────────────────────────────────────
# 메인 API
# ─────────────────────────────────────────────────────────────────────

def generate_slide_response(
    deck: dict,
    lecture_dir: Path,
    user_instruction: str,
    focus_slide: Optional[dict] = None,
    forced_layout: Optional[str] = None,
    neighbor_briefs: Optional[list] = None,
) -> dict:
    """AI 호출 → JSON 응답 파싱.

    Args:
        forced_layout: UI에서 명시적으로 선택한 layout. 있으면 강한 강제 신호로 프롬프트에 포함.

    Returns: {slide, reasoning?, speaker_note?, memo_signals?}

    ★자료는 매 슬라이드에 싣지 않는다(2026-06-24). 따라서 강의별 explicit cache(자료 본문 캐싱)도
    더 이상 쓰지 않는다 — 한 장 프롬프트는 짧아서 캐시가 불필요하다. 자료는 초안 단계 전용.
    """
    if not user_instruction.strip():
        raise ValueError("user_instruction이 비어 있습니다.")

    prompt = build_prompt(deck, lecture_dir, user_instruction, focus_slide, forced_layout,
                          neighbor_briefs)
    ai = _get_slide_ai()
    response_text = ai.process_message(prompt, history=[], images=[], execute_tool=None)

    parsed = extract_json(response_text)
    if "slide" not in parsed:
        raise ValueError(
            "AI 응답에 'slide' 키가 없습니다. 응답: " + response_text[:300]
        )

    # forced_layout이 지정됐는데 AI가 다른 걸 골랐으면 강제 교체 (안전망)
    if forced_layout:
        actual = parsed["slide"].get("layout")
        if actual != forced_layout:
            print(f"[SlideAI] forced_layout 안전망: AI가 '{actual}'을 골랐지만 '{forced_layout}'으로 강제")
            parsed["slide"]["layout"] = forced_layout

    return parsed


# ─────────────────────────────────────────────────────────────────────
# 자료 → 슬라이드 초안 목록 (일괄 생성의 1단계: outline)
# ─────────────────────────────────────────────────────────────────────

_OUTLINE_SYSTEM_PROMPT = """당신은 강의 슬라이드 기획자다. 주어진 강의 자료(책 본문 등)와 강의 메타를 읽고,
강의 한 편의 흐름에 맞는 **슬라이드 초안 목록**을 설계한다. 각 항목은 '슬라이드 한 장'을 만들기 위한 한 줄 지시문이다.

# 원칙
1. 첫 장은 표지(강의 제목·핵심 요지). 이후 도입 → 전개 → 정리의 자연스러운 흐름.
2. 한 장 = 한 명제. 각 instruction은 그 슬라이드가 말할 **단 하나의 핵심 주장**을 담는다.
   라벨('AI의 역사')이 아니라 그 슬라이드에서 무엇을 말할지가 분명해야 한다.
3. **자료에 충실** — 자료의 프레임·사실·고유명사에서 뽑고 일반 지식으로 지어내지 말 것.
4. instruction은 슬라이드 생성 AI가 그대로 받아 한 장을 만들 수 있도록 구체적으로(주제 + 각도). 한국어로.
5. 너무 잘게 쪼개지 말 것 — 분량(분)을 참고해 적정한 장수로.

# 출력 (JSON 한 객체만, 다른 텍스트 금지)
{"slides": [{"instruction": "이 슬라이드가 담을 명제/주제 (한 줄)"}, {"instruction": "..."}]}
"""


def outline_from_materials(
    deck: dict, lecture_dir: Path, count: Optional[int] = None, existing_count: int = 0
) -> list:
    """강의 자료를 읽어 슬라이드(instruction) 목록을 만든다 (슬라이드 일괄생성의 1단계).

    Returns: [{"instruction": str}, ...]
    호출자는 이 목록을 한 장씩 generate_slide_response로 돌려 강의 끝에 덧붙인다.
    existing_count > 0 이면 '이어붙이는' 생성이므로 전체 표지를 새로 만들지 않게 안내한다
    (1부·2부를 나눠 생성하는 흐름).
    """
    materials_text = _build_full_materials_text(deck, lecture_dir)
    if not materials_text.strip():
        raise ValueError("강의 자료가 없습니다. 먼저 자료(파일/메모)를 추가하세요.")

    cfg = _load_system_ai_config()
    api_key = (cfg.get("apiKey") or "").strip()
    if not api_key:
        raise RuntimeError("시스템 AI API 키가 설정되지 않았습니다. (설정 → 시스템 AI)")

    if count is not None:
        try:
            count = max(1, min(40, int(count)))
        except (TypeError, ValueError):
            count = None

    count_directive = (
        f"슬라이드를 **정확히 {count}장** 설계하라."
        if count else
        "강의 분량과 자료의 밀도에 맞춰 적정한 장수(보통 6~20장)로 설계하라."
    )
    if existing_count and existing_count > 0:
        flow_directive = (
            f"★이 강의엔 이미 슬라이드 {existing_count}장이 있고, 이번에 만드는 슬라이드는 그 **뒤에 이어붙는다**. "
            "전체 강의 표지를 새로 만들지 말고 곧장 내용 슬라이드부터 구성하라. "
            "이 자료가 새로운 부(部)의 시작이라면 짧은 부 도입(섹션 표지) 1장 정도만 허용."
        )
    else:
        flow_directive = "첫 장은 표지(강의 제목·핵심 요지)로 시작해 도입→전개→정리 흐름으로 구성하라."
    user_prompt = "\n\n".join([
        f"# 강의 메타\n{_build_meta_block(deck)}",
        f"# 지시\n{count_directive} {flow_directive} 위 원칙대로 슬라이드 목록을 JSON으로 출력하라.",
        f"# 강의 자료 (이 내용에서 직접 추출)\n{materials_text}",
    ])

    from providers import get_provider
    provider = get_provider(
        (cfg.get("provider") or "anthropic").strip(),
        api_key=api_key,
        model=(cfg.get("model") or "").strip(),
        system_prompt=_OUTLINE_SYSTEM_PROMPT,
        tools=[],
    )
    provider.init_client()
    response_text = provider.process_message(user_prompt, history=[], images=[], execute_tool=None)

    parsed = extract_json(response_text)
    slides = parsed.get("slides")
    if not isinstance(slides, list) or not slides:
        raise ValueError("AI가 슬라이드 목록을 만들지 못했습니다. 응답: " + (response_text or "")[:300])

    out = []
    for s in slides:
        if isinstance(s, dict):
            instr = (s.get("instruction") or s.get("title") or "").strip()
        else:
            instr = str(s).strip()
        if instr:
            out.append({"instruction": instr})
    if not out:
        raise ValueError("유효한 슬라이드 초안이 없습니다.")
    if count:
        out = out[:count]
    return out
