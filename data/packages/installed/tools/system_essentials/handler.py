import os
import glob
import re
import time
import fnmatch
import unicodedata
import subprocess
import shlex
import shutil
import json
from datetime import datetime
from pathlib import Path

# file_find 무경계 재귀 glob 방지 — 매 호출 홈 전체(node_modules·캐시)를 색인 없이
# stat 하던 게 타임아웃 원인. 절대-dead 가지치기 + 시간 예산으로 바운드.
# ★ 절대-dead 목록은 file_index(포식 substrate)와 *공유* — fs_query 와 같은 단일 출처라
#   드리프트 없음. path-substring 판정이라 ~/Library 통째가 아니라 캐시류만 쳐냄(iCloud 보존).
try:
    from file_index import ABSOLUTE_DEAD_SUBSTR as _DEAD_SUBSTR
except Exception:  # import 경로 미확보 시 폴백(동일 내용)
    _DEAD_SUBSTR = (
        "/System/", "/Applications/", "/Library/Caches/",
        "/Library/Application Support/", "/Library/Containers/",
        "/Library/Group Containers/", "/node_modules/", "/.Trash", ".app/",
        "/__pycache__/", "/site-packages/", "/.venv/", "/venv/",
        "/.git/", "/DerivedData/", "/.gradle/", "/.cargo/", "/.npm/",
    )
_FIND_DEADLINE_S = 25.0  # 엔진 타임아웃 전에 부분결과라도 반환


def _is_dead_dir(path):
    """절대-dead(설치트리·캐시) 디렉토리면 True — walk 가 안 들어감(의도 불문 제외)."""
    p = path.rstrip("/") + "/"
    return any(n in p for n in _DEAD_SUBSTR)


def _bounded_find(root, basename_pat, max_results):
    """root 하위를 바운드 재귀 순회 — 정크 가지치기 + dot-dir 스킵(glob ** 와 동일) + 시간 예산.

    무한정 walk 로 시스템을 멈추지 않는다. 시간 초과/상한 도달 시 partial=True 로 알린다.
    """
    deadline = time.time() + _FIND_DEADLINE_S
    # macOS 한글 파일명=NFD(자모분해), 패턴은 보통 NFC → fnmatch 바이트비교가 침묵 누락.
    # 양쪽을 NFC 로 정규화해 비교(mdfind 는 정규화하지만 fnmatch 는 안 함. forage_map #33).
    pat = unicodedata.normalize("NFC", basename_pat)
    pat_lower = pat.lower()
    matches, partial = [], False
    for dirpath, dirs, files in os.walk(root, topdown=True):
        if time.time() > deadline:
            partial = True
            break
        # 가지치기: 절대-dead(공유 목록, path-substring) + dot-dir. 제자리 수정으로 walk 가 안 들어감.
        #   ~/Library 통째가 아니라 캐시류만 → ~/Library/Mobile Documents(iCloud) 는 보존.
        dirs[:] = [d for d in dirs
                   if not d.startswith(".") and not _is_dead_dir(os.path.join(dirpath, d))]
        # 매칭: 파일 + 디렉토리 둘 다 (glob.glob 은 둘 다 매칭했음 — 예: .epub 번들·iCloud 책은 디렉토리).
        # macOS 파일시스템은 대소문자 무시 → 소문자 비교로 맞춤.
        for name in files + dirs:
            nfc = unicodedata.normalize("NFC", name)
            if fnmatch.fnmatch(nfc.lower(), pat_lower):
                matches.append(os.path.join(dirpath, name))
                if len(matches) >= max_results:
                    return matches, True
    return matches, partial

# 시스템 AI 전용 상태 폴더 (data/system_ai_state/)
DATA_PATH = Path(__file__).parent.parent.parent.parent
SYSTEM_AI_STATE_PATH = DATA_PATH / "system_ai_state"


def get_state_paths(project_path: str, agent_id: str = None) -> dict:
    """에이전트별 상태 파일 경로 반환

    - 시스템 AI: data/system_ai_state/
    - 프로젝트 에이전트: projects/{project_id}/agent_{agent_id}_*.json
    """
    project_path = Path(project_path).resolve()

    # 시스템 AI인지 확인 (project_path가 data 폴더 또는 "."인 경우)
    if str(project_path).endswith("data") or project_path == Path(".").resolve():
        state_dir = SYSTEM_AI_STATE_PATH
        prefix = ""
    else:
        # 프로젝트 에이전트는 프로젝트 폴더에 상태 저장
        state_dir = project_path
        # 에이전트별 파일명 접두사 (agent_id가 있으면 사용)
        prefix = f"agent_{agent_id}_" if agent_id else ""

    # 폴더 생성
    state_dir.mkdir(parents=True, exist_ok=True)

    return {
        "todo": state_dir / f"{prefix}todo_state.json",
        "question": state_dir / f"{prefix}question_state.json",
        "plan_mode": state_dir / f"{prefix}plan_mode_state.json",
        "plan_file": state_dir / f"{prefix}current_plan.md"
    }


# 위험한 명령어 패턴 (정규식, 사용자 승인 필요)
# 단어 경계(\b)를 사용하여 'add'에서 'dd'가 매칭되는 등의 오탐 방지
import re

DANGEROUS_PATTERNS_RE = [
    # 파일 삭제/수정
    r'\brm\s', r'\brmdir\b', r'\bunlink\b',
    # 권한 관련
    r'\bsudo\b', r'\bchmod\b', r'\bchown\b', r'\bchgrp\b',
    # 디스크 관련 위험 명령어
    r'\bdd\s', r'\bmkfs\b', r'\bformat\b', r'\bdiskutil\s+erase\b', r'\bdiskutil\s+partitionDisk\b',
    # 시스템 종료/재시작
    r'\bshutdown\b', r'\breboot\b', r'\bhalt\b',
    # 프로세스 종료
    r'\bkill\s', r'\bkillall\b', r'\bpkill\b',
]

_DANGEROUS_RE = re.compile('|'.join(DANGEROUS_PATTERNS_RE), re.IGNORECASE)


def _validate_path_in_scope(path: str, project_path: str) -> str | None:
    """경로가 프로젝트 범위 내인지 검증. 벗어나면 에러 메시지 반환, 정상이면 None

    - 절대 경로: 허용 (에이전트가 명시적으로 지정한 시스템 파일 수정 등)
    - 상대 경로가 ../로 프로젝트 밖으로 나가는 경우만 차단
    """
    # 원래 입력이 절대 경로면 허용 (시스템 파일 수정, 다른 프로젝트 파일 접근 등)
    if os.path.isabs(path):
        return None

    # 상대 경로: 프로젝트 범위 안에 있는지 확인
    abs_path = os.path.abspath(os.path.join(project_path, path))
    abs_project = os.path.abspath(project_path)
    if not abs_path.startswith(abs_project + os.sep) and abs_path != abs_project:
        return f"Error: 프로젝트 범위를 벗어나는 상대 경로입니다: {path} → {abs_path}"
    return None


def is_dangerous_command(command: str) -> bool:
    """명령어가 위험한지 검사 (정규식 단어 경계 사용)"""
    return bool(_DANGEROUS_RE.search(command))

def _get_path(tool_input: dict) -> str:
    """file_path, path, target 중 사용 가능한 경로 반환 ('~' 홈 디렉토리 확장).
    expanduser 는 '~' 없는 절대/상대 경로엔 무영향이라 기존 동작은 불변이다.
    ('~/...' 가 project_path 아래로 잘못 붙어 파일을 못 찾던 버그 방지 — read/write/edit 공통)"""
    raw = tool_input.get("file_path") or tool_input.get("path") or tool_input.get("target") or ""
    return os.path.expanduser(raw) if raw else raw


def _truthy(v) -> bool:
    """체크박스 값 해석 — bool 또는 흔한 참 표기 문자열."""
    if isinstance(v, bool):
        return v
    return str(v).strip().lower() in ("true", "yes", "on", "1", "y", "checked", "x", "v", "✓", "예", "네")


def _fill_pdf(path: Path, data: dict, output: str, flatten: bool) -> dict:
    """PDF 폼 채우기 — PyMuPDF widget. data 없으면 필드 목록 반환(introspection)."""
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    try:
        # --- introspection: 채울 수 있는 필드 열거 ---
        if not data:
            fields = []
            for page in doc:
                for w in (page.widgets() or []):
                    t = w.field_type_string
                    f = {"name": w.field_name, "type": t, "value": w.field_value}
                    if t in ("CheckBox", "RadioButton"):
                        try:
                            f["states"] = w.button_states()
                        except Exception:
                            pass
                    if t in ("ComboBox", "ListBox"):
                        f["options"] = list(getattr(w, "choice_values", None) or [])
                    fields.append(f)
            return {"success": True, "mode": "fields", "format": "pdf",
                    "items": fields, "count": len(fields), "path": str(path)}

        # --- fill ---
        filled, seen = [], set()
        for page in doc:
            for w in (page.widgets() or []):
                name = w.field_name
                if name not in data:
                    continue
                val = data[name]
                t = w.field_type_string
                if t == "CheckBox":
                    w.field_value = _truthy(val)
                else:
                    w.field_value = str(val)
                w.update()
                filled.append(name)
                seen.add(name)
        unmatched = [k for k in data if k not in seen]
        if flatten and hasattr(doc, "bake"):
            doc.bake()  # 폼 필드를 정적 내용으로 구움 → 제출용(수정 불가)
        # PyMuPDF는 열려 있는 원본에 직접 덮어쓸 수 없다 — output==template(제자리 채우기)이면
        # 임시 파일에 저장 후 교체.
        if os.path.abspath(output) == os.path.abspath(str(path)):
            tmp = output + ".tmp"
            doc.save(tmp, garbage=3, deflate=True)
            doc.close()
            os.replace(tmp, output)
            return {"success": True, "path": os.path.abspath(output), "format": "pdf",
                    "filled": filled, "unmatched": unmatched, "flattened": bool(flatten)}
        doc.save(output, garbage=3, deflate=True)
        return {"success": True, "path": os.path.abspath(output), "format": "pdf",
                "filled": filled, "unmatched": unmatched, "flattened": bool(flatten)}
    finally:
        if not getattr(doc, "is_closed", False):
            doc.close()


def _set_para_text(paragraph, text: str) -> None:
    """문단 텍스트를 통째로 교체 — 첫 run 에 넣고 나머지 run 은 비운다.
    ({{자리표시자}}가 여러 run 에 쪼개져도 문단 단위로 안전하게 치환)"""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _fill_docx(path: Path, data: dict, output: str) -> dict:
    """DOCX {{자리표시자}} 치환. data 없으면 자리표시자 이름 목록 반환."""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    ph = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

    def _paras(document):
        for p in document.paragraphs:
            yield p
        for tbl in document.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    # --- introspection ---
    if not data:
        keys = set()
        for p in _paras(doc):
            for m in ph.finditer(p.text):
                keys.add(m.group(1).strip())
        return {"success": True, "mode": "fields", "format": "docx",
                "items": [{"name": k, "type": "text"} for k in sorted(keys)],
                "count": len(keys), "path": str(path)}

    # --- fill ---
    filled = set()
    for p in _paras(doc):
        if "{{" not in p.text:
            continue
        original = p.text
        used = [k.strip() for k in ph.findall(original) if k.strip() in data]
        if not used:
            continue
        new = ph.sub(lambda m: str(data.get(m.group(1).strip(), m.group(0))), original)
        _set_para_text(p, new)
        filled.update(used)
    unmatched = [k for k in data if k not in filled]
    doc.save(output)
    return {"success": True, "path": os.path.abspath(output), "format": "docx",
            "filled": sorted(filled), "unmatched": unmatched}


def execute(tool_input: dict, context) -> str:
    """ToolContext 기반 신규 시그니처."""
    tool_name = context.tool_name
    project_path = context.project_path
    agent_id = context.agent_id

    # 단일 액션 패턴: read {format} 통합 액션. format 명시 또는 확장자 자동 인식.
    if tool_name == "read_op":
        # 파이프라인 자동 바인딩: path 가 없으면 직전 step 결과에서 파일 경로 추출.
        # "방금 찾은 파일을 읽기"([self:file_find]{...} | take: 1 >> [self:read]) 조합 개통.
        if not (tool_input.get("path") or tool_input.get("file_path") or tool_input.get("target")):
            prev = tool_input.get("_prev_result") or tool_input.get("params", {}).get("_prev_result", "")
            if prev:
                try:
                    from ibl_executors import _extract_path_from_prev
                    extracted = _extract_path_from_prev(prev if isinstance(prev, str) else json.dumps(prev, ensure_ascii=False))
                    if extracted:
                        tool_input = {**tool_input, "path": extracted}
                except Exception:  # noqa: BLE001 — 추출 실패 시 기존 경로 없음 흐름
                    pass
        fmt = (tool_input.get("format") or "").strip().lower()
        if not fmt:
            raw = tool_input.get("path") or ""
            ext = os.path.splitext(raw)[1].lower().lstrip(".")
            if ext == "pdf":
                fmt = "pdf"
            elif ext in ("docx", "doc"):
                fmt = "docx"
            elif ext in ("xlsx", "xlsm"):
                fmt = "xlsx"
            else:
                fmt = "text"
        # 분기 후 tool_name 재할당해 기존 코드로 위임
        if fmt == "pdf":
            tool_name = "read_pdf"
        elif fmt == "docx":
            tool_name = "read_docx"
        elif fmt == "xlsx":
            tool_name = "read_xlsx"
        else:
            tool_name = "read_file"

    try:
        if tool_name == "read_file":
            raw_path = _get_path(tool_input)
            # system_docs/ 경로는 어떤 프로젝트에서든 data/system_docs/로 매핑
            if raw_path.startswith("system_docs/") and not os.path.isabs(raw_path):
                from runtime_utils import get_base_path
                path = str(get_base_path() / "data" / raw_path)
            else:
                path = os.path.join(project_path, raw_path)
            # blocks: 문서를 *타입 있는 문서 IR* items 로 반환(2026-07-03 승격 — 옛 {text}
            # 문단 조각은 표면에서 마크다운이 생으로 보였음). markdown_to_blocks(backend/doc_ir)
            # 가 heading/list/quote/table/code/divider 를 살려 blocks 뷰·render_document 가
            # 그대로 소비. docx·pdf 읽기의 자체 IR 방출과 같은 통화로 3경로 정렬.
            # 원문은 message 로도 보존. 어느 파일이든 쓰는 일반 표시 옵션.
            if tool_input.get("blocks"):
                from doc_ir import markdown_to_blocks
                with open(path, 'r', encoding='utf-8') as f:
                    _txt = f.read()
                _parts = markdown_to_blocks(_txt)
                return json.dumps({"success": True, "items": _parts, "message": _txt,
                                   "path": path, "count": len(_parts)}, ensure_ascii=False)
            offset = tool_input.get("offset", 0) or 0
            limit = tool_input.get("limit")
            file_size = os.path.getsize(path)

            with open(path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            total_lines = len(lines)

            # offset/limit 적용
            if offset > 0 or limit is not None:
                end = min(offset + limit, total_lines) if limit else total_lines
                selected = lines[offset:end]
                content = ''.join(selected)
                header = f"[줄 {offset}-{min(end, total_lines)-1} / 전체 {total_lines}줄, {file_size:,}바이트]\n"
                return header + content
            else:
                # 전체 읽기 — 대용량 파일 방어 (1MB 제한)
                MAX_READ_SIZE = 1_000_000
                content = ''.join(lines)
                if len(content) > MAX_READ_SIZE:
                    content = content[:MAX_READ_SIZE]
                    content += f"\n\n... (파일이 {file_size // 1000}KB로 큽니다. 처음 1MB만 표시. offset/limit으로 부분 읽기를 사용하세요. 전체 {total_lines}줄)"
                return content

        elif tool_name == "write_file":
            raw_path = _get_path(tool_input)
            if not raw_path:
                return json.dumps({"success": False, "error": "파일 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
            path = os.path.join(project_path, raw_path)

            # 새 파일 + bare 파일명(디렉토리 없음) → outputs/ 폴더로 자동 리다이렉트
            redirected = False
            if (not os.path.isabs(raw_path)
                    and os.sep not in raw_path and '/' not in raw_path
                    and not os.path.exists(path)):
                raw_path = os.path.join("outputs", raw_path)
                path = os.path.join(project_path, raw_path)
                redirected = True

            scope_err = _validate_path_in_scope(path, project_path)
            if scope_err:
                return scope_err
            os.makedirs(os.path.dirname(path), exist_ok=True)
            content = tool_input["content"]
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            abs_path = os.path.abspath(path)
            result = {"success": True, "path": abs_path, "size": len(content)}
            if redirected:
                result["redirected_to"] = "outputs/"
            return json.dumps(result, ensure_ascii=False)

        elif tool_name == "fill_op":
            # 양식 채우기 — read 의 짝(문서→문서). PDF 폼 / DOCX 자리표시자.
            raw_path = _get_path(tool_input)
            if not raw_path:
                return json.dumps({"success": False, "error": "템플릿 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
            tpl = Path(raw_path)
            if not tpl.is_absolute():
                tpl = Path(project_path) / tpl
            if not tpl.exists():
                return json.dumps({"success": False, "error": f"템플릿을 찾을 수 없습니다: {tpl}"}, ensure_ascii=False)

            data = tool_input.get("data") or {}
            if isinstance(data, str):
                try:
                    data = json.loads(data)
                except Exception:
                    return json.dumps({"success": False, "error": "data 는 필드→값 매핑(JSON 객체)이어야 합니다."}, ensure_ascii=False)
            flatten = _truthy(tool_input.get("flatten", False))

            ext = tpl.suffix.lower().lstrip(".")

            # 출력 경로 — data 있을 때만 필요. 생략 시 원본 옆 _filled.
            out_raw = tool_input.get("output") or tool_input.get("output_path")
            if out_raw:
                out = out_raw if os.path.isabs(out_raw) else os.path.join(project_path, out_raw)
            else:
                out = str(tpl.with_name(f"{tpl.stem}_filled{tpl.suffix}"))
            if data:
                os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)

            try:
                if ext == "pdf":
                    res = _fill_pdf(tpl, data, out, flatten)
                elif ext in ("docx", "doc"):
                    res = _fill_docx(tpl, data, out)
                else:
                    return json.dumps({"success": False,
                                       "error": f"채우기를 지원하지 않는 형식입니다: .{ext} (지원: pdf, docx)"}, ensure_ascii=False)
            except Exception as e:  # noqa: BLE001
                return json.dumps({"success": False, "error": f"양식 채우기 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)
            return json.dumps(res, ensure_ascii=False)

        elif tool_name == "list_directory":
            dir_path = os.path.join(project_path, os.path.expanduser(tool_input.get("dir_path") or tool_input.get("path") or tool_input.get("target") or "."))
            items = os.listdir(dir_path)
            text = "\n".join(items)
            # === 공유 통화 table {columns, rows} (비파괴 ADD) ===
            # 파일 목록 → [이름, 크기, 수정일, 경로]. 디렉터리는 크기 "".
            rows = []
            records = []  # records 통화(보편) — 파일=명사. 선언 returns:records와 일치.
            for name in items:
                full = os.path.join(dir_path, name)
                try:
                    st = os.stat(full)
                    is_dir = os.path.isdir(full)
                    size = "" if is_dir else st.st_size
                    mtime = datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M")
                except OSError:
                    is_dir, size, mtime = False, "", ""
                abs_full = os.path.abspath(full)
                rows.append([name, size, mtime, abs_full])
                records.append({
                    "title": name + ("/" if is_dir else ""),
                    "meta": " · ".join(x for x in [
                        ("디렉터리" if is_dir else (f"{size:,}B" if isinstance(size, int) else None)),
                        (mtime or None),
                    ] if x),
                    "summary": "",
                    "url": abs_full,
                })
            table = {"columns": ["이름", "크기", "수정일", "경로"], "rows": rows}
            return json.dumps({"text": text, "table": table, "items": records}, ensure_ascii=False)

        elif tool_name == "grep_files":
            pattern = tool_input["pattern"]
            file_pattern = tool_input.get("file_pattern", "**/*")
            # 검색 루트: path > root_path > project_path (glob_files 와 동일 규칙).
            # ★과거엔 root_path 만 읽어 src 가 광고하는 path 가 조용히 무시됐다(param 불일치 버그).
            raw_root = tool_input.get("path") or tool_input.get("root_path") or "."
            expanded = os.path.expanduser(raw_root)
            root = expanded if os.path.isabs(expanded) else os.path.join(project_path, expanded)
            # 정규식이 기본이다(grep/ripgrep·Claude Grep 습관과 동일, src desc도 "정규식"으로 광고).
            # 과거엔 기본이 fixed-string('a|b' 를 리터럴로 취급)이라 alternation·메타문자가
            # 조용히 No matches 가 되는 silent-failure 함정이었다. regex=false 로 리터럴 강제 가능.
            use_regex = tool_input.get("regex", True)
            # output_mode (Claude Grep 습관): content(매칭 라인, 기본) / files_with_matches(파일 목록) /
            # count(파일별 매칭 수). 알 수 없는 값은 content 로 폴백.
            output_mode = (tool_input.get("output_mode") or "content").lower()
            if output_mode not in ("content", "files_with_matches", "count"):
                output_mode = "content"
            max_results = 100
            # 줄 수만으로는 토큰 폭발을 못 막는다 — 미니파이드/JSON 한 줄이 수만 자면
            # 100줄로도 수십만 자가 된다. 줄 길이·총량 상한을 함께 둔다.
            MAX_LINE_CHARS = 500       # 한 줄 매칭 내용 상한 (초과 시 잘라 표시)
            MAX_TOTAL_CHARS = 40_000   # content 누적 총량 상한 (초과 시 검색 중단 + 좁히기 안내)
            total_chars = 0
            hit_size_cap = False

            # 검색 제외 디렉토리/확장자 (바이너리, 캐시, VCS)
            SKIP_DIRS = {'.git', '.svn', '__pycache__', 'node_modules', '.venv', 'venv', '.tox', '.mypy_cache'}
            SKIP_EXTS = {'.pyc', '.pyo', '.so', '.dylib', '.dll', '.exe', '.bin',
                         '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp', '.svg',
                         '.mp3', '.mp4', '.wav', '.avi', '.mov', '.mkv',
                         '.zip', '.gz', '.tar', '.rar', '.7z', '.bz2',
                         '.woff', '.woff2', '.ttf', '.eot', '.otf',
                         '.pdf', '.doc', '.docx', '.xls', '.xlsx',
                         '.db', '.sqlite', '.sqlite3'}

            # 검색할 파일 목록 가져오기 (불필요 파일 필터링)
            # root 가 단일 파일이면 그 파일만 검색한다. file_pattern 기본값 '**/*' 를 파일 경로에
            # join 하면 'file.py/**/*' 가 되어 빈 목록 → 조용한 "No matches" 버그가 된다.
            # 모델(Claude Grep 습관)은 흔히 단일 파일을 겨누므로 이 케이스를 명시 처리한다.
            # 사용자가 직접 지정한 파일이므로 SKIP 확장자/디렉토리 필터도 적용하지 않는다.
            if os.path.isfile(root):
                files = [root]
            else:
                files = glob.glob(os.path.join(root, file_pattern), recursive=True)
                files = [
                    f for f in files
                    if os.path.isfile(f)
                    and os.path.splitext(f)[1].lower() not in SKIP_EXTS
                    and not any(skip in f.split(os.sep) for skip in SKIP_DIRS)
                ]

            results = []
            match_rows = []  # 공유 통화 table용 [파일, 줄번호, 내용]
            file_counts = {}  # output_mode=files_with_matches/count용 {파일: 매칭 수}
            file_order = []   # 파일 첫 등장 순서 보존
            # 잘못된 정규식(짝 안 맞는 괄호 등)은 크래시 대신 리터럴 검색으로 폴백한다 —
            # 절대 크래시도 침묵 실패도 만들지 않는다. 폴백 시 결과에 안내를 덧붙인다.
            regex_pattern = None
            regex_error = None
            if use_regex:
                try:
                    regex_pattern = re.compile(pattern)
                except re.error as e:
                    regex_error = str(e)
                    use_regex = False  # 리터럴로 폴백
            search_done = False

            for file_path in files:
                if search_done:
                    break
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        for line_num, line in enumerate(f, 1):
                            if use_regex:
                                matched = regex_pattern.search(line)
                            else:
                                matched = pattern in line

                            if matched:
                                rel_path = os.path.relpath(file_path, project_path)
                                snippet = line.rstrip()
                                if len(snippet) > MAX_LINE_CHARS:
                                    snippet = snippet[:MAX_LINE_CHARS] + " …(줄 잘림)"
                                results.append(f"{rel_path}:{line_num}: {snippet}")
                                match_rows.append([rel_path, line_num, snippet])
                                if rel_path not in file_counts:
                                    file_order.append(rel_path)
                                file_counts[rel_path] = file_counts.get(rel_path, 0) + 1
                                total_chars += len(snippet)
                                if len(results) >= max_results or total_chars >= MAX_TOTAL_CHARS:
                                    search_done = True
                                    hit_size_cap = total_chars >= MAX_TOTAL_CHARS
                                    break
                except (UnicodeDecodeError, PermissionError, OSError):
                    continue

            regex_note = ""
            if regex_error:
                regex_note = (f"\n(주의: 정규식 '{pattern}' 컴파일 실패 [{regex_error}] → "
                              f"리터럴 문자열로 검색했습니다. 의도가 정규식이면 패턴을 고치세요.)")

            if not results:
                return f"No matches found for: {pattern}{regex_note}"

            if search_done and hit_size_cap:
                truncated = ("\n... (결과가 너무 큽니다 — root_path/file_pattern 으로 범위를 좁히거나, "
                             "output_mode='files_with_matches'/'count' 로 먼저 분포를 보세요.)")
            elif search_done:
                truncated = "\n... (결과가 {0}개에 도달하여 검색 중단)".format(max_results)
            else:
                truncated = ""
            truncated += regex_note

            if output_mode == "files_with_matches":
                # 매칭된 파일 목록만 (라인 무관). 단일 통화 items(행 dict).
                text = "\n".join(file_order) + truncated
                items = [{"파일": fp} for fp in file_order]
                return json.dumps({"text": text, "items": items}, ensure_ascii=False)

            if output_mode == "count":
                # 파일별 매칭 수. 단일 통화 items(행 dict).
                rows = [[fp, file_counts[fp]] for fp in file_order]
                text = "\n".join(f"{fp}: {cnt}" for fp, cnt in rows) + truncated
                items = [{"파일": fp, "매칭 수": cnt} for fp, cnt in rows]
                return json.dumps({"text": text, "items": items}, ensure_ascii=False)

            # output_mode == "content" (기본): 매칭 라인 전체.
            text = "\n".join(results) + truncated
            # === 단일 통화 items(행 dict — 파일/줄번호/내용) ===
            items = [{"파일": r[0], "줄번호": r[1], "내용": r[2]} for r in match_rows]
            return json.dumps({"text": text, "items": items}, ensure_ascii=False)

        elif tool_name == "get_current_time":
            fmt = tool_input.get("format", "%Y-%m-%d %H:%M:%S")
            return datetime.now().strftime(fmt)

        elif tool_name == "ai_ask":
            # 시스템 AI 원샷 호출 — 도구·다단계 없이 경량 LLM 으로 즉답. [self:ask]
            # 능력=어휘: 앱(선언형/커스텀)이 raw fetch 없이 IBL 로 AI 를 부른다.
            prompt = (tool_input.get("prompt") or "").strip()
            if not prompt:
                return json.dumps({"success": False, "error": "prompt(지시/질문)가 필요합니다."}, ensure_ascii=False)
            # context 명시가 없으면 파이프 입력(_prev_result)을 맥락으로 받는다 → 조합 가능
            # (예: [sense:search_gnews]{...} >> [self:ask]{prompt: "요약해줘"}).
            ctx = tool_input.get("context")
            if ctx is None:
                prev = tool_input.get("_prev_result")
                if prev not in (None, ""):
                    ctx = prev if isinstance(prev, str) else json.dumps(prev, ensure_ascii=False)
            if ctx is not None and not isinstance(ctx, str):
                ctx = json.dumps(ctx, ensure_ascii=False)
            message = f"{ctx}\n\n---\n\n{prompt}" if ctx else prompt
            sys_prompt = (tool_input.get("system") or
                          "당신은 앱에 내장된 유능한 조수입니다. 사용자의 지시에 간결하고 정확하게 답하세요. "
                          "불필요한 서론·맺음말 없이 요청한 결과만 반환하세요.")
            try:
                from consciousness_agent import lightweight_ai_call
                answer = lightweight_ai_call(message, system_prompt=sys_prompt, role="background")
            except Exception as e:
                return json.dumps({"success": False, "error": f"AI 호출 실패: {e}"}, ensure_ascii=False)
            if not answer:
                return json.dumps({"success": False, "error": "AI 응답을 받지 못했습니다(모델 미설정 가능)."}, ensure_ascii=False)
            return json.dumps({"result": answer, "text": answer}, ensure_ascii=False)

        elif tool_name == "glob_files":
            pattern = tool_input["pattern"]

            # 검색 루트 결정 (우선순위: path > root_path > project_path)
            # - 절대경로(/...): 그대로 사용 → 컴퓨터 어디든 검색 가능
            # - ~ 시작: 홈 디렉토리로 확장
            # - 상대경로: project_path 기준
            # - 미지정: project_path
            raw_root = tool_input.get("path") or tool_input.get("root_path") or "."
            expanded = os.path.expanduser(raw_root)
            if os.path.isabs(expanded):
                root = expanded
            else:
                root = os.path.join(project_path, expanded)

            try:
                max_results = int(tool_input.get("max_results", 200))
            except (TypeError, ValueError):
                max_results = 200

            partial = False
            if "/" not in pattern and "**" not in pattern:
                # 재귀 basename 검색(지배적 케이스) — 바운드 walk: 정크 가지치기 + 시간예산.
                # glob.glob('~/**/*X*') 의 색인없는 홈 전체 stat(=타임아웃)을 회피.
                matches, partial = _bounded_find(root, pattern, max_results)
            else:
                # 명시 경로/패턴(`/`·`**` 포함) — 보통 앵커돼 빠름. glob 유지.
                search_pattern = pattern if "**" in pattern else f"**/{pattern}"
                try:
                    matches = glob.glob(os.path.join(root, search_pattern), recursive=True)
                except Exception as e:
                    return f"검색 오류: {e}"

            absolute_paths = sorted(os.path.abspath(m) for m in matches)
            total = len(absolute_paths)
            truncated = (total > max_results > 0) or partial
            if total > max_results > 0:
                absolute_paths = absolute_paths[:max_results]

            if absolute_paths:
                header_parts = [f"{total}개 매칭"]
                if partial:
                    header_parts.append(f"(시간예산 {int(_FIND_DEADLINE_S)}초/상한 도달 — 부분 결과. 더 좁은 path 로 재검색하거나 fs_query(OS색인) 사용 권장)")
                elif truncated:
                    header_parts.append(f"(상위 {max_results}개만 반환 — 더 많으면 max_results 또는 더 좁은 path 사용)")
                header_parts.append(f"root: {root}")
                header = " | ".join(header_parts)
                text = header + "\n" + "\n".join(absolute_paths)
                # === 공유 통화 table {columns, rows} (비파괴 ADD) ===
                # 매칭 파일 → [이름, 크기, 수정일, 경로]. 디렉터리는 크기 "".
                rows = []
                records = []  # records 통화(보편) — 파일=명사. 선언 returns:records와 일치.
                for p in absolute_paths:
                    try:
                        is_dir = os.path.isdir(p)
                        st = os.stat(p)
                        size = "" if is_dir else st.st_size
                        mtime = datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M")
                    except OSError:
                        is_dir, size, mtime = False, "", ""
                    rows.append([os.path.basename(p), size, mtime, p])
                    records.append({
                        "title": os.path.basename(p) + ("/" if is_dir else ""),
                        "meta": " · ".join(x for x in [
                            ("디렉터리" if is_dir else (f"{size:,}B" if isinstance(size, int) else None)),
                            (mtime or None),
                        ] if x),
                        "summary": "",
                        "url": p,
                    })
                table = {"columns": ["이름", "크기", "수정일", "경로"], "rows": rows}
                return json.dumps({"text": text, "table": table, "items": records}, ensure_ascii=False)

            # 결과 없을 때 — 안내 메시지에 path 옵션 힌트 포함
            hint = (
                f"매칭 없음: pattern={pattern!r} root={root}\n"
                "힌트: 프로젝트 밖을 검색하려면 path 파라미터를 사용하세요. "
                '예: {pattern: "*.docx", path: "~/Desktop"} 또는 {pattern: "*.docx", path: "/Users"}'
            )
            return hint

        elif tool_name == "edit_file":
            file_path = os.path.join(project_path, _get_path(tool_input))
            scope_err = _validate_path_in_scope(file_path, project_path)
            if scope_err:
                return scope_err
            old_string = tool_input["old_string"]
            new_string = tool_input["new_string"]

            # 파일 읽기
            if not os.path.exists(file_path):
                return f"Error: 파일이 존재하지 않습니다: {_get_path(tool_input)}"

            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # old_string이 파일에 있는지 확인
            count = content.count(old_string)
            if count == 0:
                return f"Error: 교체할 문자열을 찾을 수 없습니다. 파일 내용을 다시 확인하세요."
            elif count > 1:
                return f"Error: 교체할 문자열이 {count}번 발견되었습니다. 더 구체적인 문자열을 지정하세요."

            # 교체 수행
            new_content = content.replace(old_string, new_string, 1)

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(new_content)

            # 절대 경로로 반환 (에이전트 간 경로 혼동 방지)
            return f"Successfully edited {os.path.abspath(file_path)}"

        elif tool_name == "run_command":
            command = tool_input.get("command", "").strip()
            timeout = min(tool_input.get("timeout", 60), 300)  # 최대 300초
            approved = tool_input.get("approved", False)

            if not command:
                return "Error: 명령어가 비어있습니다."

            # 위험한 명령어 감지 - 승인되지 않았으면 승인 요청
            if not approved and is_dangerous_command(command):
                return f"__REQUIRES_APPROVAL__:{command}"

            # 명령어 실행
            try:
                result = subprocess.run(
                    command,
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    cwd=project_path
                )

                output = ""
                if result.stdout:
                    output += result.stdout
                if result.stderr:
                    output += f"\n[stderr]\n{result.stderr}" if output else result.stderr

                if result.returncode != 0:
                    output += f"\n[exit code: {result.returncode}]"

                return output.strip() if output else "(명령어가 출력 없이 완료됨)"

            except subprocess.TimeoutExpired:
                return f"Error: 명령어 실행 시간 초과 ({timeout}초)"
            except Exception as e:
                return f"Error: 명령어 실행 실패: {e}"

        elif tool_name == "copy_path":
            _src = tool_input.get("src") or tool_input.get("source")  # src 우선(코퍼스/자연어), source 별칭
            _dst = tool_input.get("dest") or tool_input.get("destination")
            if not _src or not _dst:
                return "Error: src(원본)와 dest(대상) 경로가 필요합니다."
            src = os.path.join(project_path, os.path.expanduser(_src))
            dst = os.path.join(project_path, os.path.expanduser(_dst))
            scope_err = _validate_path_in_scope(dst, project_path)
            if scope_err:
                return scope_err

            if not os.path.exists(src):
                return f"Error: 원본이 존재하지 않습니다: {src}"

            # 대상 상위 디렉토리 생성
            os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)

            if os.path.isdir(src):
                # 폴더 복사 (대상이 이미 있으면 삭제 후 복사)
                if os.path.exists(dst):
                    shutil.rmtree(dst)
                shutil.copytree(src, dst)
                count = sum(len(files) for _, _, files in os.walk(dst))
                return f"폴더를 복사했습니다: {os.path.abspath(dst)} ({count}개 파일)"
            else:
                # 파일 복사
                shutil.copy2(src, dst)
                return f"파일을 복사했습니다: {os.path.abspath(dst)}"

        elif tool_name == "move_path":
            _src = tool_input.get("src") or tool_input.get("source")  # src 우선(코퍼스/자연어), source 별칭
            _dst = tool_input.get("dest") or tool_input.get("destination")
            if not _src or not _dst:
                return "Error: src(원본)와 dest(대상) 경로가 필요합니다."
            src = os.path.join(project_path, os.path.expanduser(_src))
            dst = os.path.join(project_path, os.path.expanduser(_dst))
            scope_err = _validate_path_in_scope(dst, project_path)
            if scope_err:
                return scope_err

            if not os.path.exists(src):
                return f"Error: 원본이 존재하지 않습니다: {src}"

            # 대상 상위 디렉토리 생성
            os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)

            # 대상이 이미 있으면 삭제
            if os.path.exists(dst):
                if os.path.isdir(dst):
                    shutil.rmtree(dst)
                else:
                    os.remove(dst)

            shutil.move(src, dst)
            return f"이동 완료: {os.path.abspath(dst)}"

        elif tool_name == "delete_path":
            target = os.path.join(project_path, os.path.expanduser(tool_input["path"]))
            scope_err = _validate_path_in_scope(target, project_path)
            if scope_err:
                return scope_err

            if not os.path.exists(target):
                return f"Error: 경로가 존재하지 않습니다: {target}"

            abs_target = os.path.abspath(target)

            if os.path.isdir(target):
                count = sum(len(files) for _, _, files in os.walk(target))
                shutil.rmtree(target)
                return f"폴더를 삭제했습니다: {abs_target} ({count}개 파일 포함)"
            else:
                os.remove(target)
                return f"파일을 삭제했습니다: {abs_target}"

        elif tool_name == "make_directory":
            raw_path = _get_path(tool_input)
            if not raw_path:
                return json.dumps({"success": False, "error": "폴더 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
            target = os.path.join(project_path, os.path.expanduser(raw_path))
            scope_err = _validate_path_in_scope(target, project_path)
            if scope_err:
                return scope_err
            abs_target = os.path.abspath(target)
            if os.path.isfile(abs_target):
                return json.dumps({"success": False, "error": f"같은 이름의 파일이 이미 있습니다: {abs_target}"}, ensure_ascii=False)
            existed = os.path.isdir(abs_target)
            os.makedirs(abs_target, exist_ok=True)
            return json.dumps({"success": True, "path": abs_target, "existed": existed}, ensure_ascii=False)

        elif tool_name == "read_pdf":
            import fitz  # PyMuPDF
            file_path = tool_input.get("file_path") or tool_input.get("path")  # path 별칭 수용(read 일관성)
            pages = tool_input.get("pages")

            if not file_path:
                return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

            path = Path(file_path)
            if not path.is_absolute():
                path = Path(project_path) / path

            if not path.exists():
                return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

            try:
                doc = fitz.open(str(path))
                metadata = doc.metadata
                total_pages = doc.page_count

                if pages is not None and isinstance(pages, list):
                    target_pages = [p for p in pages if 0 <= p < total_pages]
                else:
                    target_pages = range(total_pages)

                extracted_text = ""
                for pno in target_pages:
                    page = doc.load_page(pno)
                    extracted_text += f"\n--- Page {pno + 1} ---\n"
                    extracted_text += page.get_text()

                doc.close()

                # 문서 IR blocks(비파괴) — pdf 텍스트를 문단 블록으로. read(x.pdf) >> document.
                pdf_blocks = []
                for para in extracted_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        pdf_blocks.append({"type": "paragraph", "text": para})
                res = {
                    "success": True,
                    "metadata": metadata,
                    "total_pages": total_pages,
                    "extracted_pages_count": len(list(target_pages)),
                    "text": extracted_text,
                    "blocks": pdf_blocks or [{"type": "paragraph", "text": extracted_text}],
                }
                return json.dumps(res, ensure_ascii=False)

            except Exception as e:
                return json.dumps({"success": False, "error": f"PDF를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)

        elif tool_name == "read_docx":
            from docx import Document as DocxDocument
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            import zipfile

            file_path = tool_input.get("file_path") or tool_input.get("path")
            extract_images = tool_input.get("extract_images", True)

            # 부분 읽기 파라미터 — 큰 docx의 컨텍스트 잠식 방지
            # 블록 = 문단(p) 또는 표(tbl) 하나
            try:
                offset = max(0, int(tool_input.get("offset", 0) or 0))
            except (TypeError, ValueError):
                offset = 0
            limit_raw = tool_input.get("limit")
            try:
                limit = int(limit_raw) if limit_raw is not None else None
            except (TypeError, ValueError):
                limit = None
            try:
                max_blocks = int(tool_input.get("max_blocks", 300))
            except (TypeError, ValueError):
                max_blocks = 300

            if not file_path:
                return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

            path = Path(file_path)
            if not path.is_absolute():
                path = Path(project_path) / path

            if not path.exists():
                return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

            try:
                doc = DocxDocument(str(path))

                # --- 이미지 추출 ---
                images_info = []
                images_dir = None
                if extract_images:
                    # ZIP에서 word/media/ 추출
                    with zipfile.ZipFile(str(path), 'r') as zf:
                        media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
                        if media_files:
                            images_dir = path.parent / f"{path.stem}_images"
                            images_dir.mkdir(exist_ok=True)
                            for mf in media_files:
                                img_name = os.path.basename(mf)
                                img_data = zf.read(mf)
                                img_path = images_dir / img_name
                                with open(img_path, 'wb') as f:
                                    f.write(img_data)
                                images_info.append({
                                    "name": img_name,
                                    "saved_path": str(img_path),
                                    "size": len(img_data),
                                })

                # --- 이미지-문단 관계 매핑 ---
                # relationship ID → 이미지 파일명 매핑
                rid_to_image = {}
                try:
                    for rel in doc.part.rels.values():
                        if "image" in rel.reltype:
                            rid_to_image[rel.rId] = os.path.basename(rel.target_ref)
                except Exception:
                    pass

                # --- 텍스트 추출 (이미지 위치 마커 포함) ---
                from docx.oxml.ns import qn
                extracted_parts = []
                # === 공유 통화 문서 IR blocks (비파괴 ADD) ===
                # 단락→{type:paragraph,text}, Heading 스타일→{type:heading,level,text},
                # 표→{type:table,columns,rows}. read(보고서.docx) >> document{pdf} 포맷변환용.
                blocks = []
                para_index = 0

                def _heading_level(el):
                    # w:pPr/w:pStyle 의 val 이 Heading N / 제목 N 이면 레벨 반환, 아니면 None
                    try:
                        ppr = el.find(qn('w:pPr'))
                        if ppr is None:
                            return None
                        pstyle = ppr.find(qn('w:pStyle'))
                        if pstyle is None:
                            return None
                        val = (pstyle.get(qn('w:val')) or "")
                    except Exception:
                        return None
                    low = val.lower()
                    # "Heading1", "Heading 1", "Title", "제목 1" 등 흡수
                    if low in ("title",) or "title" in low:
                        return 1
                    m = re.search(r'(\d+)', val)
                    if "heading" in low or "제목" in val:
                        if m:
                            return min(6, max(1, int(m.group(1))))
                        return 1
                    return None

                for element in doc.element.body:
                    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

                    if tag == 'p':  # 문단
                        para_text = element.text or ""
                        # 하위 요소에서 텍스트 수집
                        runs_text = []
                        for run in element.iter(qn('w:r')):
                            for t in run.iter(qn('w:t')):
                                if t.text:
                                    runs_text.append(t.text)
                        full_text = ''.join(runs_text) if runs_text else para_text

                        # 인라인 이미지 탐지
                        for drawing in element.iter(qn('w:drawing')):
                            for blip in drawing.iter(qn('a:blip')):
                                embed_id = blip.get(qn('r:embed'))
                                if embed_id and embed_id in rid_to_image:
                                    img_name = rid_to_image[embed_id]
                                    full_text += f"\n[이미지: {img_name}]"

                        if full_text.strip():
                            txt = full_text.strip()
                            extracted_parts.append(txt)
                            level = _heading_level(element)
                            if level is not None:
                                blocks.append({"type": "heading", "level": level, "text": txt})
                            else:
                                blocks.append({"type": "paragraph", "text": txt})
                        para_index += 1

                    elif tag == 'tbl':  # 표
                        table_lines = ["[표 시작]"]
                        tbl_rows = []
                        for row in element.iter(qn('w:tr')):
                            cells = []
                            for cell in row.iter(qn('w:tc')):
                                cell_texts = []
                                for p in cell.iter(qn('w:p')):
                                    t_parts = []
                                    for t in p.iter(qn('w:t')):
                                        if t.text:
                                            t_parts.append(t.text)
                                    cell_texts.append(''.join(t_parts))
                                cells.append(' '.join(cell_texts))
                            table_lines.append(" | ".join(cells))
                            tbl_rows.append(cells)
                        table_lines.append("[표 끝]")
                        extracted_parts.append('\n'.join(table_lines))
                        # 표 블록: 첫 행을 columns, 나머지를 rows (길이 정규화)
                        if tbl_rows:
                            tcols = tbl_rows[0]
                            ncol = len(tcols)
                            tbody = []
                            for r in tbl_rows[1:]:
                                rr = list(r)
                                if len(rr) < ncol:
                                    rr += [""] * (ncol - len(rr))
                                elif len(rr) > ncol:
                                    rr = rr[:ncol]
                                tbody.append(rr)
                            blocks.append({"type": "table", "columns": tcols, "rows": tbody})

                # --- 부분 읽기 슬라이싱 ---
                total_blocks = len(extracted_parts)

                # 적용 한계 결정: 명시적 limit > max_blocks(기본 안전망)
                # max_blocks=0이면 무제한
                if limit is not None and limit >= 0:
                    effective_limit = limit
                elif max_blocks > 0:
                    effective_limit = max_blocks
                else:
                    effective_limit = None  # 무제한

                start = min(offset, total_blocks)
                if effective_limit is None:
                    end = total_blocks
                else:
                    end = min(start + effective_limit, total_blocks)

                sliced = extracted_parts[start:end]
                returned_blocks = len(sliced)
                truncated = end < total_blocks

                text = '\n\n'.join(sliced)

                metadata = {
                    "filename": path.name,
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "total_images": len(images_info),
                    "total_blocks": total_blocks,
                    "offset": start,
                    "returned_blocks": returned_blocks,
                    "truncated": truncated,
                }
                if truncated:
                    metadata["next_offset"] = end
                    metadata["hint"] = f"전체 {total_blocks}블록 중 {start}~{end-1}만 반환됨. 다음 호출에 offset={end}로 이어 읽기."
                if images_dir:
                    metadata["images_dir"] = str(images_dir)

                res = {
                    "success": True,
                    "metadata": metadata,
                    "text": text,
                }
                # 문서 IR blocks 통화 (비파괴 ADD) — 전체 문서 구조.
                # 부분 읽기로 text가 잘려도 blocks 자체는 문서 IR이라 전체 제공.
                if blocks:
                    res["blocks"] = blocks
                if images_info:
                    res["images"] = images_info

                return json.dumps(res, ensure_ascii=False)

            except Exception as e:
                return json.dumps({"success": False, "error": f"DOCX를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)

        elif tool_name == "read_xlsx":
            import openpyxl

            file_path = tool_input.get("file_path") or tool_input.get("path")
            sheet_name = tool_input.get("sheet")  # 특정 시트만 (생략 시 전체)
            try:
                max_rows = int(tool_input.get("max_rows", 200) or 200)
            except (TypeError, ValueError):
                max_rows = 200

            if not file_path:
                return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

            path = Path(file_path)
            if not path.is_absolute():
                path = Path(project_path) / path
            if not path.exists():
                return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

            try:
                # read_only=True(대용량 안전), data_only=True(수식 대신 계산값)
                wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                all_sheets = list(wb.sheetnames)
                targets = [sheet_name] if sheet_name else all_sheets

                parts = []
                for sn in targets:
                    if sn not in all_sheets:
                        parts.append(f"### 시트: {sn} — 없음")
                        continue
                    ws = wb[sn]
                    rows_text = []
                    truncated = False
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i >= max_rows:
                            truncated = True
                            break
                        cells = ["" if c is None else str(c) for c in row]
                        rows_text.append("\t".join(cells))
                    header = f"### 시트: {sn} ({ws.max_row}행 × {ws.max_column}열)"
                    if truncated:
                        header += f" — 처음 {max_rows}행만 (max_rows로 조정)"
                    parts.append(header + "\n" + "\n".join(rows_text))

                # === 공유 통화 table {columns, rows} (비파괴 ADD) ===
                # 첫 행을 헤더로. 시트 미지정이면 가장 큰(셀 수 최대) 시트, 지정이면 그 시트.
                # read(데이터.xlsx) >> chart / spreadsheet 로 흐를 수 있게 함.
                def _num(v):
                    # 가능하면 숫자로 — table 통화는 값이 숫자면 더 유용
                    if isinstance(v, (int, float)):
                        return v
                    if isinstance(v, str):
                        s = v.strip().replace(",", "")
                        try:
                            return int(s)
                        except (ValueError, TypeError):
                            pass
                        try:
                            return float(s)
                        except (ValueError, TypeError):
                            pass
                    return "" if v is None else v

                table = None
                if sheet_name and sheet_name in all_sheets:
                    table_sheets = [sheet_name]
                else:
                    # 가장 큰 시트 선택 (max_row * max_column)
                    table_sheets = sorted(
                        [sn for sn in all_sheets],
                        key=lambda sn: (wb[sn].max_row or 0) * (wb[sn].max_column or 0),
                        reverse=True,
                    )[:1]
                if table_sheets:
                    tsn = table_sheets[0]
                    tws = wb[tsn]
                    all_rows = []
                    for i, row in enumerate(tws.iter_rows(values_only=True)):
                        if i >= max_rows:
                            break
                        all_rows.append(list(row))
                    if all_rows:
                        columns = ["" if c is None else str(c) for c in all_rows[0]]
                        body = []
                        for r in all_rows[1:]:
                            # columns 길이에 맞춰 패딩/절단
                            cells = [_num(c) for c in r]
                            if len(cells) < len(columns):
                                cells += [""] * (len(columns) - len(cells))
                            elif len(cells) > len(columns):
                                cells = cells[:len(columns)]
                            body.append(cells)
                        table = {"columns": columns, "rows": body}

                wb.close()

                res = {
                    "success": True,
                    "sheet_count": len(all_sheets),
                    "sheets": all_sheets,
                    "text": "\n\n".join(parts),
                }
                if table is not None:
                    res["table"] = table
                return json.dumps(res, ensure_ascii=False)

            except Exception as e:
                return json.dumps({"success": False, "error": f"XLSX를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)

        elif tool_name == "spreadsheet":
            # [table:spreadsheet] — 행 데이터 → xlsx 산출 (값만, 수식/서식은 범위 밖)
            import openpyxl

            # >> 파이프: 이전 액션 결과(_prev_result)에 table 통화가 있으면 자동 수용
            if not tool_input.get("table") and not tool_input.get("rows") and not tool_input.get("sheets"):
                _pr = tool_input.get("_prev_result")
                if _pr:
                    try:
                        _po = json.loads(_pr) if isinstance(_pr, str) else _pr
                        _rows_src = _po.get("items") if isinstance(_po, dict) else None
                        if isinstance(_po, dict) and isinstance(_po.get("table"), dict) and _po["table"].get("rows"):
                            tool_input["table"] = _po["table"]
                        elif isinstance(_rows_src, list) and _rows_src:
                            # 단일 통화 items(행 dict) → table 투영: 목록형 생산자도 엑셀로.
                            _it = [x for x in _rows_src if isinstance(x, dict)]
                            if _it and "title" in _it[0]:
                                # records-관습 카드 → 사람친화 4열
                                tool_input["table"] = {
                                    "columns": ["제목", "정보", "요약", "링크"],
                                    "rows": [[x.get("title", ""), x.get("meta", ""),
                                              x.get("summary", ""), x.get("url", "")] for x in _it],
                                }
                            elif _it:
                                # 임의 items(world_bank 연도/지표 등) → 키 순서=열 generic 재구성
                                _cols = list(_it[0].keys())
                                tool_input["table"] = {"columns": _cols,
                                                       "rows": [[x.get(c) for c in _cols] for x in _it]}
                    except Exception:
                        pass
            # 표준 테이블 통화 수용: {columns, rows} → headers/rows (table:chart와 동일 통화).
            # 같은 통화 한 벌이 차트로도, 표로도 흘러감 (데이터 소스 >> 시각화/표).
            _table = tool_input.get("table")
            if isinstance(_table, dict) and not tool_input.get("sheets") and not tool_input.get("rows"):
                if _table.get("columns") and not tool_input.get("headers"):
                    tool_input["headers"] = _table["columns"]
                if _table.get("rows"):
                    tool_input["rows"] = _table["rows"]

            raw_path = _get_path(tool_input)
            if not raw_path:
                return json.dumps({"success": False, "error": "출력 파일 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
            if not raw_path.lower().endswith((".xlsx", ".xlsm")):
                raw_path += ".xlsx"
            path = os.path.join(project_path, raw_path)

            # write_file과 동일: bare 파일명 → outputs/ 리다이렉트
            redirected = False
            if (not os.path.isabs(raw_path)
                    and os.sep not in raw_path and '/' not in raw_path
                    and not os.path.exists(path)):
                raw_path = os.path.join("outputs", raw_path)
                path = os.path.join(project_path, raw_path)
                redirected = True

            scope_err = _validate_path_in_scope(path, project_path)
            if scope_err:
                return scope_err

            def _coerce(cell):
                # openpyxl이 받는 타입(str/int/float/bool/None)으로 강제
                if cell is None or isinstance(cell, (str, int, float, bool)):
                    return cell
                return str(cell)

            def _fill(ws, rows):
                for r in (rows or []):
                    if isinstance(r, (list, tuple)):
                        ws.append([_coerce(c) for c in r])
                    else:
                        ws.append([_coerce(r)])

            try:
                sheets = tool_input.get("sheets")
                wb = openpyxl.Workbook()
                if sheets and isinstance(sheets, dict):
                    # 다중 시트: {시트명: rows}
                    first = True
                    for sn, srows in sheets.items():
                        ws = wb.active if first else wb.create_sheet()
                        ws.title = str(sn)[:31]  # xlsx 시트명 최대 31자
                        _fill(ws, srows)
                        first = False
                    sheet_summary = list(sheets.keys())
                else:
                    # 단일 시트: rows (+ 선택 headers)
                    ws = wb.active
                    ws.title = str(tool_input.get("sheet_name", "Sheet1"))[:31]
                    headers = tool_input.get("headers")
                    if headers:
                        ws.append([_coerce(c) for c in headers])
                    _fill(ws, tool_input.get("rows"))
                    sheet_summary = [ws.title]

                os.makedirs(os.path.dirname(path), exist_ok=True)
                wb.save(path)
                abs_path = os.path.abspath(path)
                result = {"success": True, "path": abs_path, "sheets": sheet_summary}
                if redirected:
                    result["redirected_to"] = "outputs/"
                return json.dumps(result, ensure_ascii=False)

            except Exception as e:
                return json.dumps({"success": False, "error": f"XLSX를 쓰는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)

        elif tool_name == "todo_write":
            todos = tool_input.get("todos", [])
            paths = get_state_paths(project_path, agent_id)

            # 상태 저장
            state = {
                "todos": todos,
                "updated_at": datetime.now().isoformat()
            }

            with open(paths["todo"], 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=2)

            # 결과 포맷팅 (텍스트 기반, Anthropic 스타일)
            result_lines = ["Todo list updated:"]
            for i, todo in enumerate(todos, 1):
                status = todo["status"]
                if status == "in_progress":
                    result_lines.append(f"  {i}. [in_progress] {todo['activeForm']}")
                elif status == "completed":
                    result_lines.append(f"  {i}. [completed] {todo['content']}")
                else:
                    result_lines.append(f"  {i}. [pending] {todo['content']}")

            return "\n".join(result_lines)

        elif tool_name == "ask_user_question":
            # GoalEval 재실행 중에는 사용자에게 질문할 수 없음
            try:
                from thread_context import get_current_task_id
                current_task = get_current_task_id() or ""
                if current_task.startswith("goal_retry_"):
                    return (
                        "현재 평가 재실행 중이므로 사용자에게 질문할 수 없습니다. "
                        "보유한 정보만으로 최선의 답변을 작성하세요. "
                        "확실하지 않은 부분은 가정을 명시하고 진행하세요."
                    )
            except Exception:
                pass

            questions = tool_input.get("questions", [])
            paths = get_state_paths(project_path, agent_id)

            # 질문 상태 저장 (프론트엔드에서 폴링)
            state = {
                "questions": questions,
                "status": "pending",  # pending, answered
                "answers": None,
                "created_at": datetime.now().isoformat()
            }

            with open(paths["question"], 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=2)

            # 특수 마커와 함께 반환 - 프론트엔드가 질문 UI를 표시
            return "[[QUESTION_PENDING]]사용자에게 질문을 전달했습니다. 응답을 기다리는 중..."

        elif tool_name == "enter_plan_mode":
            paths = get_state_paths(project_path, agent_id)

            # 계획 모드 상태 저장
            state = {
                "active": True,
                "phase": "exploring",  # exploring, designing, reviewing, finalizing
                "entered_at": datetime.now().isoformat(),
                "plan_file": str(paths["plan_file"])
            }

            with open(paths["plan_mode"], 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=2)

            # 계획 파일 초기화
            with open(paths["plan_file"], 'w', encoding='utf-8') as f:
                f.write(f"# 구현 계획\n\n생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n## 목표\n\n(작성 중...)\n\n## 구현 단계\n\n1. \n\n## 수정할 파일\n\n- \n\n## 테스트 방법\n\n- \n")

            return "[[PLAN_MODE_ENTERED]]계획 모드로 진입했습니다. 코드를 탐색하고 구현 계획을 수립한 후 exit_plan_mode를 호출하세요."

        elif tool_name == "exit_plan_mode":
            paths = get_state_paths(project_path, agent_id)

            # 계획 모드 상태 확인
            if not paths["plan_mode"].exists():
                return "Error: 계획 모드가 활성화되어 있지 않습니다."

            with open(paths["plan_mode"], 'r', encoding='utf-8') as f:
                state = json.load(f)

            if not state.get("active"):
                return "Error: 계획 모드가 활성화되어 있지 않습니다."

            # 계획 파일 읽기
            plan_content = ""
            if paths["plan_file"].exists():
                with open(paths["plan_file"], 'r', encoding='utf-8') as f:
                    plan_content = f.read()

            # 상태 업데이트 - 승인 대기
            state["phase"] = "awaiting_approval"
            state["plan_content"] = plan_content

            with open(paths["plan_mode"], 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=2)

            return f"[[PLAN_APPROVAL_REQUESTED]]계획 수립이 완료되었습니다. 사용자 승인을 기다리는 중...\n\n---\n{plan_content}"

        else:
            return f"Unknown tool: {tool_name}"

    except Exception as e:
        return f"Error: {str(e)}"
