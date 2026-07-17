"""system_essentials 오피스/문서 op 층 (2026-07-18 모듈화 — 1500줄 규칙)

handler.py 에서 verbatim 이동: fill(PDF 폼/DOCX 자리표시자)·read_pdf/read_docx/
read_xlsx·spreadsheet(xlsx 산출) 분기 몸통 + 전용 헬퍼(_fill_pdf/_fill_docx/_truthy).
handler 가 real-estate 선례(load_module spec-load)로 붙여 위임한다.
_get_path 는 여기가 정의처 — handler 가 별칭으로 재수출(다른 분기도 계속 사용).
"""
import os
import re
import json
from pathlib import Path

def _get_path(tool_input: dict) -> str:
    """file_path, path, target 중 사용 가능한 경로 반환 ('~' 홈 디렉토리 확장).
    expanduser 는 '~' 없는 절대/상대 경로엔 무영향이라 기존 동작은 불변이다.
    ('~/...' 가 project_path 아래로 잘못 붙어 파일을 못 찾던 버그 방지 — read/write/edit 공통)"""
    raw = tool_input.get("file_path") or tool_input.get("path") or tool_input.get("target") or ""
    return os.path.expanduser(raw) if raw else raw


def _truthy(v) -> bool:
    """체크박스 값 해석 — bool 또는 흔한 참 표기 문자열."""
    if isinstance(v, bool):
        return v
    return str(v).strip().lower() in ("true", "yes", "on", "1", "y", "checked", "x", "v", "✓", "예", "네")


def _fill_pdf(path: Path, data: dict, output: str, flatten: bool) -> dict:
    """PDF 폼 채우기 — PyMuPDF widget. data 없으면 필드 목록 반환(introspection)."""
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    try:
        # --- introspection: 채울 수 있는 필드 열거 ---
        if not data:
            fields = []
            for page in doc:
                for w in (page.widgets() or []):
                    t = w.field_type_string
                    f = {"name": w.field_name, "type": t, "value": w.field_value}
                    if t in ("CheckBox", "RadioButton"):
                        try:
                            f["states"] = w.button_states()
                        except Exception:
                            pass
                    if t in ("ComboBox", "ListBox"):
                        f["options"] = list(getattr(w, "choice_values", None) or [])
                    fields.append(f)
            return {"success": True, "mode": "fields", "format": "pdf",
                    "items": fields, "count": len(fields), "path": str(path)}

        # --- fill ---
        filled, seen = [], set()
        for page in doc:
            for w in (page.widgets() or []):
                name = w.field_name
                if name not in data:
                    continue
                val = data[name]
                t = w.field_type_string
                if t == "CheckBox":
                    w.field_value = _truthy(val)
                else:
                    w.field_value = str(val)
                w.update()
                filled.append(name)
                seen.add(name)
        unmatched = [k for k in data if k not in seen]
        if flatten and hasattr(doc, "bake"):
            doc.bake()  # 폼 필드를 정적 내용으로 구움 → 제출용(수정 불가)
        # PyMuPDF는 열려 있는 원본에 직접 덮어쓸 수 없다 — output==template(제자리 채우기)이면
        # 임시 파일에 저장 후 교체.
        if os.path.abspath(output) == os.path.abspath(str(path)):
            tmp = output + ".tmp"
            doc.save(tmp, garbage=3, deflate=True)
            doc.close()
            os.replace(tmp, output)
            return {"success": True, "path": os.path.abspath(output), "format": "pdf",
                    "filled": filled, "unmatched": unmatched, "flattened": bool(flatten)}
        doc.save(output, garbage=3, deflate=True)
        return {"success": True, "path": os.path.abspath(output), "format": "pdf",
                "filled": filled, "unmatched": unmatched, "flattened": bool(flatten)}
    finally:
        if not getattr(doc, "is_closed", False):
            doc.close()


def _set_para_text(paragraph, text: str) -> None:
    """문단 텍스트를 통째로 교체 — 첫 run 에 넣고 나머지 run 은 비운다.
    ({{자리표시자}}가 여러 run 에 쪼개져도 문단 단위로 안전하게 치환)"""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _fill_docx(path: Path, data: dict, output: str) -> dict:
    """DOCX {{자리표시자}} 치환. data 없으면 자리표시자 이름 목록 반환."""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    ph = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

    def _paras(document):
        for p in document.paragraphs:
            yield p
        for tbl in document.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    # --- introspection ---
    if not data:
        keys = set()
        for p in _paras(doc):
            for m in ph.finditer(p.text):
                keys.add(m.group(1).strip())
        return {"success": True, "mode": "fields", "format": "docx",
                "items": [{"name": k, "type": "text"} for k in sorted(keys)],
                "count": len(keys), "path": str(path)}

    # --- fill ---
    filled = set()
    for p in _paras(doc):
        if "{{" not in p.text:
            continue
        original = p.text
        used = [k.strip() for k in ph.findall(original) if k.strip() in data]
        if not used:
            continue
        new = ph.sub(lambda m: str(data.get(m.group(1).strip(), m.group(0))), original)
        _set_para_text(p, new)
        filled.update(used)
    unmatched = [k for k in data if k not in filled]
    doc.save(output)
    return {"success": True, "path": os.path.abspath(output), "format": "docx",
            "filled": sorted(filled), "unmatched": unmatched}


def fill_op(tool_input: dict, project_path: str) -> str:
    # 양식 채우기 — read 의 짝(문서→문서). PDF 폼 / DOCX 자리표시자.
    raw_path = _get_path(tool_input)
    if not raw_path:
        return json.dumps({"success": False, "error": "템플릿 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
    tpl = Path(raw_path)
    if not tpl.is_absolute():
        tpl = Path(project_path) / tpl
    if not tpl.exists():
        return json.dumps({"success": False, "error": f"템플릿을 찾을 수 없습니다: {tpl}"}, ensure_ascii=False)

    data = tool_input.get("data") or {}
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception:
            return json.dumps({"success": False, "error": "data 는 필드→값 매핑(JSON 객체)이어야 합니다."}, ensure_ascii=False)
    flatten = _truthy(tool_input.get("flatten", False))

    ext = tpl.suffix.lower().lstrip(".")

    # 출력 경로 — data 있을 때만 필요. 생략 시 원본 옆 _filled.
    out_raw = tool_input.get("output") or tool_input.get("output_path")
    if out_raw:
        out = out_raw if os.path.isabs(out_raw) else os.path.join(project_path, out_raw)
    else:
        out = str(tpl.with_name(f"{tpl.stem}_filled{tpl.suffix}"))
    if data:
        os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)

    try:
        if ext == "pdf":
            res = _fill_pdf(tpl, data, out, flatten)
        elif ext in ("docx", "doc"):
            res = _fill_docx(tpl, data, out)
        else:
            return json.dumps({"success": False,
                               "error": f"채우기를 지원하지 않는 형식입니다: .{ext} (지원: pdf, docx)"}, ensure_ascii=False)
    except Exception as e:  # noqa: BLE001
        return json.dumps({"success": False, "error": f"양식 채우기 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)
    return json.dumps(res, ensure_ascii=False)


def read_pdf(tool_input: dict, project_path: str) -> str:
    import fitz  # PyMuPDF
    file_path = tool_input.get("file_path") or tool_input.get("path")  # path 별칭 수용(read 일관성)
    pages = tool_input.get("pages")

    if not file_path:
        return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

    path = Path(file_path)
    if not path.is_absolute():
        path = Path(project_path) / path

    if not path.exists():
        return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

    try:
        doc = fitz.open(str(path))
        metadata = doc.metadata
        total_pages = doc.page_count

        if pages is not None and isinstance(pages, list):
            target_pages = [p for p in pages if 0 <= p < total_pages]
        else:
            target_pages = range(total_pages)

        extracted_text = ""
        for pno in target_pages:
            page = doc.load_page(pno)
            extracted_text += f"\n--- Page {pno + 1} ---\n"
            extracted_text += page.get_text()

        doc.close()

        # 문서 IR blocks(비파괴) — pdf 텍스트를 문단 블록으로. read(x.pdf) >> document.
        pdf_blocks = []
        for para in extracted_text.split("\n\n"):
            para = para.strip()
            if para:
                pdf_blocks.append({"type": "paragraph", "text": para})
        res = {
            "success": True,
            "metadata": metadata,
            "total_pages": total_pages,
            "extracted_pages_count": len(list(target_pages)),
            "text": extracted_text,
            "blocks": pdf_blocks or [{"type": "paragraph", "text": extracted_text}],
        }
        return json.dumps(res, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"success": False, "error": f"PDF를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)


def read_docx(tool_input: dict, project_path: str) -> str:
    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import zipfile

    file_path = tool_input.get("file_path") or tool_input.get("path")
    extract_images = tool_input.get("extract_images", True)

    # 부분 읽기 파라미터 — 큰 docx의 컨텍스트 잠식 방지
    # 블록 = 문단(p) 또는 표(tbl) 하나
    try:
        offset = max(0, int(tool_input.get("offset", 0) or 0))
    except (TypeError, ValueError):
        offset = 0
    limit_raw = tool_input.get("limit")
    try:
        limit = int(limit_raw) if limit_raw is not None else None
    except (TypeError, ValueError):
        limit = None
    try:
        max_blocks = int(tool_input.get("max_blocks", 300))
    except (TypeError, ValueError):
        max_blocks = 300

    if not file_path:
        return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

    path = Path(file_path)
    if not path.is_absolute():
        path = Path(project_path) / path

    if not path.exists():
        return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

    try:
        doc = DocxDocument(str(path))

        # --- 이미지 추출 ---
        images_info = []
        images_dir = None
        if extract_images:
            # ZIP에서 word/media/ 추출
            with zipfile.ZipFile(str(path), 'r') as zf:
                media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
                if media_files:
                    images_dir = path.parent / f"{path.stem}_images"
                    images_dir.mkdir(exist_ok=True)
                    for mf in media_files:
                        img_name = os.path.basename(mf)
                        img_data = zf.read(mf)
                        img_path = images_dir / img_name
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        images_info.append({
                            "name": img_name,
                            "saved_path": str(img_path),
                            "size": len(img_data),
                        })

        # --- 이미지-문단 관계 매핑 ---
        # relationship ID → 이미지 파일명 매핑
        rid_to_image = {}
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    rid_to_image[rel.rId] = os.path.basename(rel.target_ref)
        except Exception:
            pass

        # --- 텍스트 추출 (이미지 위치 마커 포함) ---
        from docx.oxml.ns import qn
        extracted_parts = []
        # === 공유 통화 문서 IR blocks (비파괴 ADD) ===
        # 단락→{type:paragraph,text}, Heading 스타일→{type:heading,level,text},
        # 표→{type:table,columns,rows}. read(보고서.docx) >> document{pdf} 포맷변환용.
        blocks = []
        para_index = 0

        def _heading_level(el):
            # w:pPr/w:pStyle 의 val 이 Heading N / 제목 N 이면 레벨 반환, 아니면 None
            try:
                ppr = el.find(qn('w:pPr'))
                if ppr is None:
                    return None
                pstyle = ppr.find(qn('w:pStyle'))
                if pstyle is None:
                    return None
                val = (pstyle.get(qn('w:val')) or "")
            except Exception:
                return None
            low = val.lower()
            # "Heading1", "Heading 1", "Title", "제목 1" 등 흡수
            if low in ("title",) or "title" in low:
                return 1
            m = re.search(r'(\d+)', val)
            if "heading" in low or "제목" in val:
                if m:
                    return min(6, max(1, int(m.group(1))))
                return 1
            return None

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':  # 문단
                para_text = element.text or ""
                # 하위 요소에서 텍스트 수집
                runs_text = []
                for run in element.iter(qn('w:r')):
                    for t in run.iter(qn('w:t')):
                        if t.text:
                            runs_text.append(t.text)
                full_text = ''.join(runs_text) if runs_text else para_text

                # 인라인 이미지 탐지
                for drawing in element.iter(qn('w:drawing')):
                    for blip in drawing.iter(qn('a:blip')):
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id and embed_id in rid_to_image:
                            img_name = rid_to_image[embed_id]
                            full_text += f"\n[이미지: {img_name}]"

                if full_text.strip():
                    txt = full_text.strip()
                    extracted_parts.append(txt)
                    level = _heading_level(element)
                    if level is not None:
                        blocks.append({"type": "heading", "level": level, "text": txt})
                    else:
                        blocks.append({"type": "paragraph", "text": txt})
                para_index += 1

            elif tag == 'tbl':  # 표
                table_lines = ["[표 시작]"]
                tbl_rows = []
                for row in element.iter(qn('w:tr')):
                    cells = []
                    for cell in row.iter(qn('w:tc')):
                        cell_texts = []
                        for p in cell.iter(qn('w:p')):
                            t_parts = []
                            for t in p.iter(qn('w:t')):
                                if t.text:
                                    t_parts.append(t.text)
                            cell_texts.append(''.join(t_parts))
                        cells.append(' '.join(cell_texts))
                    table_lines.append(" | ".join(cells))
                    tbl_rows.append(cells)
                table_lines.append("[표 끝]")
                extracted_parts.append('\n'.join(table_lines))
                # 표 블록: 첫 행을 columns, 나머지를 rows (길이 정규화)
                if tbl_rows:
                    tcols = tbl_rows[0]
                    ncol = len(tcols)
                    tbody = []
                    for r in tbl_rows[1:]:
                        rr = list(r)
                        if len(rr) < ncol:
                            rr += [""] * (ncol - len(rr))
                        elif len(rr) > ncol:
                            rr = rr[:ncol]
                        tbody.append(rr)
                    blocks.append({"type": "table", "columns": tcols, "rows": tbody})

        # --- 부분 읽기 슬라이싱 ---
        total_blocks = len(extracted_parts)

        # 적용 한계 결정: 명시적 limit > max_blocks(기본 안전망)
        # max_blocks=0이면 무제한
        if limit is not None and limit >= 0:
            effective_limit = limit
        elif max_blocks > 0:
            effective_limit = max_blocks
        else:
            effective_limit = None  # 무제한

        start = min(offset, total_blocks)
        if effective_limit is None:
            end = total_blocks
        else:
            end = min(start + effective_limit, total_blocks)

        sliced = extracted_parts[start:end]
        returned_blocks = len(sliced)
        truncated = end < total_blocks

        text = '\n\n'.join(sliced)

        metadata = {
            "filename": path.name,
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_images": len(images_info),
            "total_blocks": total_blocks,
            "offset": start,
            "returned_blocks": returned_blocks,
            "truncated": truncated,
        }
        if truncated:
            metadata["next_offset"] = end
            metadata["hint"] = f"전체 {total_blocks}블록 중 {start}~{end-1}만 반환됨. 다음 호출에 offset={end}로 이어 읽기."
        if images_dir:
            metadata["images_dir"] = str(images_dir)

        res = {
            "success": True,
            "metadata": metadata,
            "text": text,
        }
        # 문서 IR blocks 통화 (비파괴 ADD) — 전체 문서 구조.
        # 부분 읽기로 text가 잘려도 blocks 자체는 문서 IR이라 전체 제공.
        if blocks:
            res["blocks"] = blocks
        if images_info:
            res["images"] = images_info

        return json.dumps(res, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"success": False, "error": f"DOCX를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)


def read_xlsx(tool_input: dict, project_path: str) -> str:
    import openpyxl

    file_path = tool_input.get("file_path") or tool_input.get("path")
    sheet_name = tool_input.get("sheet")  # 특정 시트만 (생략 시 전체)
    try:
        max_rows = int(tool_input.get("max_rows", 200) or 200)
    except (TypeError, ValueError):
        max_rows = 200

    if not file_path:
        return json.dumps({"success": False, "error": "file_path가 제공되지 않았습니다."}, ensure_ascii=False)

    path = Path(file_path)
    if not path.is_absolute():
        path = Path(project_path) / path
    if not path.exists():
        return json.dumps({"success": False, "error": f"파일을 찾을 수 없습니다: {path}"}, ensure_ascii=False)

    try:
        # read_only=True(대용량 안전), data_only=True(수식 대신 계산값)
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        all_sheets = list(wb.sheetnames)
        targets = [sheet_name] if sheet_name else all_sheets

        parts = []
        for sn in targets:
            if sn not in all_sheets:
                parts.append(f"### 시트: {sn} — 없음")
                continue
            ws = wb[sn]
            rows_text = []
            truncated = False
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    truncated = True
                    break
                cells = ["" if c is None else str(c) for c in row]
                rows_text.append("\t".join(cells))
            header = f"### 시트: {sn} ({ws.max_row}행 × {ws.max_column}열)"
            if truncated:
                header += f" — 처음 {max_rows}행만 (max_rows로 조정)"
            parts.append(header + "\n" + "\n".join(rows_text))

        # === 공유 통화 table {columns, rows} (비파괴 ADD) ===
        # 첫 행을 헤더로. 시트 미지정이면 가장 큰(셀 수 최대) 시트, 지정이면 그 시트.
        # read(데이터.xlsx) >> chart / spreadsheet 로 흐를 수 있게 함.
        def _num(v):
            # 가능하면 숫자로 — table 통화는 값이 숫자면 더 유용
            if isinstance(v, (int, float)):
                return v
            if isinstance(v, str):
                s = v.strip().replace(",", "")
                try:
                    return int(s)
                except (ValueError, TypeError):
                    pass
                try:
                    return float(s)
                except (ValueError, TypeError):
                    pass
            return "" if v is None else v

        table = None
        if sheet_name and sheet_name in all_sheets:
            table_sheets = [sheet_name]
        else:
            # 가장 큰 시트 선택 (max_row * max_column)
            table_sheets = sorted(
                [sn for sn in all_sheets],
                key=lambda sn: (wb[sn].max_row or 0) * (wb[sn].max_column or 0),
                reverse=True,
            )[:1]
        if table_sheets:
            tsn = table_sheets[0]
            tws = wb[tsn]
            all_rows = []
            for i, row in enumerate(tws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                all_rows.append(list(row))
            if all_rows:
                columns = ["" if c is None else str(c) for c in all_rows[0]]
                body = []
                for r in all_rows[1:]:
                    # columns 길이에 맞춰 패딩/절단
                    cells = [_num(c) for c in r]
                    if len(cells) < len(columns):
                        cells += [""] * (len(columns) - len(cells))
                    elif len(cells) > len(columns):
                        cells = cells[:len(columns)]
                    body.append(cells)
                table = {"columns": columns, "rows": body}

        wb.close()

        res = {
            "success": True,
            "sheet_count": len(all_sheets),
            "sheets": all_sheets,
            "text": "\n\n".join(parts),
        }
        if table is not None:
            res["table"] = table
        return json.dumps(res, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"success": False, "error": f"XLSX를 읽는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)


def spreadsheet(tool_input: dict, project_path: str, validate_path_in_scope) -> str:
    # [table:spreadsheet] — 행 데이터 → xlsx 산출 (값만, 수식/서식은 범위 밖)
    import openpyxl

    # >> 파이프: 이전 액션 결과(_prev_result)에 table 통화가 있으면 자동 수용
    if not tool_input.get("table") and not tool_input.get("rows") and not tool_input.get("sheets"):
        _pr = tool_input.get("_prev_result")
        if _pr:
            try:
                _po = json.loads(_pr) if isinstance(_pr, str) else _pr
                _rows_src = _po.get("items") if isinstance(_po, dict) else None
                if isinstance(_po, dict) and isinstance(_po.get("table"), dict) and _po["table"].get("rows"):
                    tool_input["table"] = _po["table"]
                elif isinstance(_rows_src, list) and _rows_src:
                    # 단일 통화 items(행 dict) → table 투영: 목록형 생산자도 엑셀로.
                    _it = [x for x in _rows_src if isinstance(x, dict)]
                    if _it and "title" in _it[0]:
                        # records-관습 카드 → 사람친화 4열
                        tool_input["table"] = {
                            "columns": ["제목", "정보", "요약", "링크"],
                            "rows": [[x.get("title", ""), x.get("meta", ""),
                                      x.get("summary", ""), x.get("url", "")] for x in _it],
                        }
                    elif _it:
                        # 임의 items(world_bank 연도/지표 등) → 키 순서=열 generic 재구성
                        _cols = list(_it[0].keys())
                        tool_input["table"] = {"columns": _cols,
                                               "rows": [[x.get(c) for c in _cols] for x in _it]}
            except Exception:
                pass
    # 표준 테이블 통화 수용: {columns, rows} → headers/rows (table:chart와 동일 통화).
    # 같은 통화 한 벌이 차트로도, 표로도 흘러감 (데이터 소스 >> 시각화/표).
    _table = tool_input.get("table")
    if isinstance(_table, dict) and not tool_input.get("sheets") and not tool_input.get("rows"):
        if _table.get("columns") and not tool_input.get("headers"):
            tool_input["headers"] = _table["columns"]
        if _table.get("rows"):
            tool_input["rows"] = _table["rows"]

    raw_path = _get_path(tool_input)
    if not raw_path:
        return json.dumps({"success": False, "error": "출력 파일 경로(path)가 지정되지 않았습니다."}, ensure_ascii=False)
    if not raw_path.lower().endswith((".xlsx", ".xlsm")):
        raw_path += ".xlsx"
    path = os.path.join(project_path, raw_path)

    # write_file과 동일: bare 파일명 → outputs/ 리다이렉트
    redirected = False
    if (not os.path.isabs(raw_path)
            and os.sep not in raw_path and '/' not in raw_path
            and not os.path.exists(path)):
        raw_path = os.path.join("outputs", raw_path)
        path = os.path.join(project_path, raw_path)
        redirected = True

    scope_err = validate_path_in_scope(path, project_path)
    if scope_err:
        return scope_err

    def _coerce(cell):
        # openpyxl이 받는 타입(str/int/float/bool/None)으로 강제
        if cell is None or isinstance(cell, (str, int, float, bool)):
            return cell
        return str(cell)

    def _fill(ws, rows):
        for r in (rows or []):
            if isinstance(r, (list, tuple)):
                ws.append([_coerce(c) for c in r])
            else:
                ws.append([_coerce(r)])

    try:
        sheets = tool_input.get("sheets")
        wb = openpyxl.Workbook()
        if sheets and isinstance(sheets, dict):
            # 다중 시트: {시트명: rows}
            first = True
            for sn, srows in sheets.items():
                ws = wb.active if first else wb.create_sheet()
                ws.title = str(sn)[:31]  # xlsx 시트명 최대 31자
                _fill(ws, srows)
                first = False
            sheet_summary = list(sheets.keys())
        else:
            # 단일 시트: rows (+ 선택 headers)
            ws = wb.active
            ws.title = str(tool_input.get("sheet_name", "Sheet1"))[:31]
            headers = tool_input.get("headers")
            if headers:
                ws.append([_coerce(c) for c in headers])
            _fill(ws, tool_input.get("rows"))
            sheet_summary = [ws.title]

        os.makedirs(os.path.dirname(path), exist_ok=True)
        wb.save(path)
        abs_path = os.path.abspath(path)
        result = {"success": True, "path": abs_path, "sheets": sheet_summary}
        if redirected:
            result["redirected_to"] = "outputs/"
        return json.dumps(result, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"success": False, "error": f"XLSX를 쓰는 중 문제가 발생했습니다: {str(e)}"}, ensure_ascii=False)
